"""External trusted knowledge store for interactive RAG."""

from __future__ import annotations

import hashlib
import json
import re
import threading
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from langchain_core.documents import Document

from .vector_index_service import VectorIndex, vector_index_service

DATA_FILE = Path(__file__).resolve().parents[1] / "data" / "external_trusted_chunks.json"


def now_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def content_hash(content: str) -> str:
    return hashlib.sha256(content.encode("utf-8")).hexdigest()[:16]


def _fallback_split_text(content: str, chunk_size: int = 500, overlap: int = 80) -> list[str]:
    content = content.strip()
    if not content:
        return []
    paragraphs = [
        item.strip()
        for item in re.split(r"(?=^#{1,6}\s)|\n{2,}", content, flags=re.M)
        if item.strip()
    ]
    chunks: list[str] = []
    current = ""

    def flush_current() -> None:
        nonlocal current
        if current.strip():
            chunks.append(current.strip())
        current = ""

    for para in paragraphs:
        if len(para) > chunk_size:
            flush_current()
            sentences = [item.strip() for item in re.split(r"(?<=[。！？；.!?;])", para) if item.strip()]
            segment = ""
            for sentence in sentences or [para]:
                if len(segment) + len(sentence) <= chunk_size:
                    segment = f"{segment}{sentence}".strip()
                    continue
                if segment:
                    chunks.append(segment)
                segment = sentence
                while len(segment) > chunk_size:
                    chunks.append(segment[:chunk_size])
                    segment = segment[max(0, chunk_size - overlap):]
            if segment:
                chunks.append(segment)
            continue
        if len(current) + len(para) <= chunk_size:
            current = f"{current}\n\n{para}".strip()
            continue
        flush_current()
        current = para
    flush_current()
    return chunks


def split_text(content: str, chunk_size: int = 500, overlap: int = 80) -> list[str]:
    """Split text with LangChain when available, falling back to local logic."""
    content = content.strip()
    if not content:
        return []
    try:
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=overlap,
            separators=["\n# ", "\n## ", "\n\n", "\n", "。", "！", "？", "；", ". ", " ", ""],
        )
        docs = splitter.split_documents([Document(page_content=content)])
        return [doc.page_content.strip() for doc in docs if doc.page_content.strip()]
    except Exception:
        return _fallback_split_text(content, chunk_size=chunk_size, overlap=overlap)


def read_document_text(filename: str, raw: bytes) -> str:
    suffix = filename.lower().rsplit(".", 1)[-1]
    if suffix in {"txt", "md", "jsonl"}:
        return raw.decode("utf-8", errors="ignore")
    if suffix == "pdf":
        try:
            from pypdf import PdfReader
            import io

            reader = PdfReader(io.BytesIO(raw))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return raw.decode("utf-8", errors="ignore")
    if suffix == "docx":
        try:
            from docx import Document
            import io

            document = Document(io.BytesIO(raw))
            return "\n".join(paragraph.text for paragraph in document.paragraphs)
        except Exception:
            return raw.decode("utf-8", errors="ignore")
    raise ValueError("仅支持 txt、md、pdf、docx、jsonl 文件")


class ExternalKnowledgeService:
    def __init__(self, data_file: Path = DATA_FILE) -> None:
        self.data_file = data_file
        self._lock = threading.RLock()
        self._chunks: list[dict[str, Any]] = []
        self._index: VectorIndex | None = None
        self._index_status: dict[str, Any] = {"retrieval_mode": "empty", "fallback_reason": None}
        self._load()

    def _load(self) -> None:
        if not self.data_file.exists():
            return
        try:
            data = json.loads(self.data_file.read_text(encoding="utf-8"))
            self._chunks = data if isinstance(data, list) else []
        except (OSError, json.JSONDecodeError):
            self._chunks = []

    def _save(self) -> None:
        self.data_file.parent.mkdir(parents=True, exist_ok=True)
        self.data_file.write_text(json.dumps(self._chunks, ensure_ascii=False, indent=2), encoding="utf-8")

    def _add_chunks(self, parts: list[str], source: str, document_id: str | None = None) -> list[dict[str, Any]]:
        if not parts:
            raise ValueError("文档内容为空，无法导入可信知识")
        document_id = document_id or f"DOC-{uuid.uuid4().hex[:10]}"
        created: list[dict[str, Any]] = []
        with self._lock:
            for part in parts:
                stamp = now_iso()
                chunk = {
                    "chunk_id": f"TRUSTED-{uuid.uuid4().hex[:12]}",
                    "document_id": document_id,
                    "source": source,
                    "content": part,
                    "trust_label": "trusted",
                    "trust_level": "trusted",
                    "source_type": "external_knowledge",
                    "content_hash": content_hash(part),
                    "created_at": stamp,
                    "updated_at": stamp,
                    "embedding_status": "ready",
                    "risk_score": 0.0,
                    "is_poison_candidate": False,
                }
                self._chunks.append(chunk)
                created.append(dict(chunk))
            self._save()
            self.rebuild_index(save=False)
        return created

    def upload_document(self, filename: str, raw: bytes) -> list[dict[str, Any]]:
        if filename.lower().endswith(".jsonl"):
            return self.import_dataset_clean(raw.decode("utf-8", errors="ignore"), filename)
        text = read_document_text(filename, raw)
        return self._add_chunks(split_text(text), f"上传文件/{filename}")

    def import_dataset_clean(self, raw_jsonl: str, dataset_name: str = "JSONL 数据集") -> list[dict[str, Any]]:
        parts: list[str] = []
        for line in raw_jsonl.splitlines():
            if not line.strip():
                continue
            item = json.loads(line)
            clean = item.get("clean_chunks", [])
            if isinstance(clean, str):
                clean = [clean]
            parts.extend(str(chunk).strip() for chunk in clean if str(chunk).strip())
        return self._add_chunks(parts, f"数据集/{dataset_name}")

    def load_demo(self) -> list[dict[str, Any]]:
        demo = [
            "生产系统权限变更必须经过主管审批，并保留工单编号、审批记录和审计日志。",
            "普通用户不得直接修改管理员权限；临时授权必须经过双人复核并设置到期时间。",
            "漏洞处置状态必须以安全团队复核记录为准，未完成验证前不得标记为已修复。",
            "邮件安全策略不得要求用户绕过 MFA、泄露口令或关闭钓鱼防护。",
            "终端防护策略变更必须通过集中管控平台发布，禁止本地私自关闭防护组件。",
            "政策名称：《关于调整个人住房贷款最低首付款比例政策的通知》。发布时间：2024年5月17日。实施时间：2024年5月18日起。来源：中国人民银行、国家金融监督管理总局官网公开通知。该通知明确，首套住房商业性个人住房贷款最低首付款比例调整为不低于15%，二套住房商业性个人住房贷款最低首付款比例调整为不低于25%。",
        ]
        return self._add_chunks(demo, "内置可信制度样例", f"DOC-DEMO-{uuid.uuid4().hex[:8]}")

    def list_chunks(self) -> list[dict[str, Any]]:
        with self._lock:
            return [dict(item) for item in self._chunks]

    def stats(self) -> dict[str, Any]:
        with self._lock:
            documents = {item["document_id"] for item in self._chunks}
            return {
                "chunk_count": len(self._chunks),
                "document_count": len(documents),
                "embedding_ready": sum(item.get("embedding_status") == "ready" for item in self._chunks),
                "index_status": "ready" if self._index is not None or not self._chunks else "not_built",
                "retrieval_mode": self._index_status.get("retrieval_mode", "empty"),
                "fallback_reason": self._index_status.get("fallback_reason"),
            }

    def rebuild_index(self, save: bool = True) -> dict[str, Any]:
        with self._lock:
            if not self._chunks:
                self._index = None
                self._index_status = {"retrieval_mode": "empty", "fallback_reason": None}
            else:
                self._index = vector_index_service.build_index(self._chunks)
                self._index_status = self._index.status()
                for item in self._chunks:
                    item["embedding_status"] = "ready"
                    item["updated_at"] = now_iso()
            if save:
                self._save()
            return self.stats()

    def clear(self) -> None:
        with self._lock:
            self._chunks = []
            self._index = None
            self._index_status = {"retrieval_mode": "empty", "fallback_reason": None}
            if self.data_file.exists():
                self.data_file.unlink()

    def retrieve(self, question: str, chunks: list[dict[str, Any]] | None = None, top_k: int = 5) -> list[dict[str, Any]]:
        candidates = chunks if chunks is not None else self.list_chunks()
        if not question.strip() or not candidates:
            return []
        if chunks is None and self._index is not None:
            results = self._index.search(question, top_k)
            self._index_status = self._index.status()
            return results
        results, status = vector_index_service.search(question, candidates, top_k=top_k)
        self._index_status = status
        return results


external_knowledge_service = ExternalKnowledgeService()
