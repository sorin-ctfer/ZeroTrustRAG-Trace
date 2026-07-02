# -*- coding: utf-8 -*-
"""Generate a Hard-dataset version of the chapter 3 document using python-docx."""
import csv, json
from collections import Counter
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent.parent
DATASET_DIR = ROOT / "实验数据集" / "mabzt_comm_dataset_hard"
RESULTS_DIR = ROOT / "实验结果" / "hard"
OUT = ROOT / "多Agent协同第三章_hard.docx"

def load_json(p): 
    with open(p, encoding="utf-8") as f: return json.load(f)
def load_csv(p): 
    with open(p, encoding="utf-8-sig", newline="") as f: return list(csv.DictReader(f))

manifest = load_json(DATASET_DIR / "manifest.json")
stats = load_json(RESULTS_DIR / "dataset_statistics.json")
run_summary = load_json(RESULTS_DIR / "run_summary.json")
validation_summary = load_csv(RESULTS_DIR / "validation_summary.csv")
risk_summary = load_csv(RESULTS_DIR / "risk_detection_summary.csv")
baseline = load_csv(RESULTS_DIR / "baseline_comparison.csv")
latency = load_csv(RESULTS_DIR / "latency_results.csv")
consensus_summary = load_csv(RESULTS_DIR / "consensus_decision_summary.csv")
risk_scores = load_csv(RESULTS_DIR / "risk_scores.csv")
agents = load_json(DATASET_DIR / "agents.json")
comm_events = load_json(DATASET_DIR / "comm_events.json")
claims = load_json(DATASET_DIR / "claims.json")
evidence = load_json(DATASET_DIR / "evidence.json")

role_counts = Counter(a.get("role") for a in agents)
attack_counts = Counter(a.get("attack_type") for a in agents)
scenario_counts = Counter(e.get("scenario") for e in comm_events)
validation_failed = int(stats.get("validation_failed", 0))
validation_passed = int(stats.get("validation_passed", 0))
validation_total = validation_failed + validation_passed

doc = Document()

def h1(text): doc.add_heading(text, level=1)
def h2(text): doc.add_heading(text, level=2)
def p(text): doc.add_paragraph(text)
def bullet(text): doc.add_paragraph(text, style='List Bullet')

def tbl(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers), style='Light Grid Accent 1')
    for i, h in enumerate(headers): table.rows[0].cells[i].text = str(h)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx+1].cells[c_idx].text = str(val)
    doc.add_paragraph()

# ── Title ──
doc.add_heading("多Agent协同通信与零信任共识实验报告（Hard 数据集）", level=0)
p(f"数据集 ID: {manifest.get('dataset_id')}  |  生成时间: {manifest.get('generated_at')}")
p(f"结果目录: {RESULTS_DIR}")

# ── 1. 数据集规模 ──
h1("一、数据集规模")
tbl(["指标", "数值"], [[k, v] for k, v in manifest.get("scale", {}).items()])
p(f"攻击家族: {', '.join(manifest.get('attack_families', []))}")

# ── 2. Agent 构成 ──
h1("二、Agent 构成")
h2("2.1 攻击类型分布")
tbl(["攻击类型", "数量"], [[k, v] for k, v in sorted(attack_counts.items())])
h2("2.2 角色分布")
tbl(["角色", "数量"], [[k, v] for k, v in sorted(role_counts.items())])
h2("2.3 信任度分布")
trust_by_type = {}
for a in agents:
    trust_by_type.setdefault(a["attack_type"], []).append(a["trust_prior"])
rows = []
for at, trusts in sorted(trust_by_type.items()):
    rows.append([at, f"{min(trusts)}-{max(trusts)}", f"{sum(trusts)/len(trusts):.1f}"])
tbl(["攻击类型", "信任度范围", "平均值"], rows)

# ── 3. 场景分布 ──
h1("三、声明场景分布")
tbl(["场景", "数量"], [[k, v] for k, v in sorted(scenario_counts.items())])

# ── 4. 零信任验证 ──
h1("四、零信任验证结果")
p(f"总声明包: {validation_total}")
p(f"通过: {validation_passed} ({validation_passed/validation_total*100:.1f}%)")
p(f"失败: {validation_failed} ({validation_failed/validation_total*100:.1f}%)")
h2("4.1 各校验项失败统计")
check_rows = []
for r in validation_summary:
    check_rows.append([r["check"], r["failed"], f"{float(r['failure_rate'])*100:.2f}%"])
tbl(["校验项", "失败数", "失败率"], check_rows)

# ── 5. BSS 风险评分 ──
h1("五、BSS 风险评分")
h2("5.1 Top 10 高风险 Agent")
top10 = sorted(risk_scores, key=lambda r: float(r.get("bss") or 0), reverse=True)[:10]
tbl(["Agent", "角色", "BSS", "状态", "根因", "真实类型"],
    [[r["agent_id"], r.get("role",""), f"{float(r.get('bss',0)):.4f}", r.get("status",""), r.get("root_cause",""), r.get("ground_truth","")] for r in top10])
h2("5.2 按真实类型汇总")
tbl(["真实类型", "Agent数", "平均BSS", "受限/隔离", "正常/观察"],
    [[r["ground_truth"], r["agent_count"], r.get("avg_bss",""), r["restricted_or_isolated"], r["normal_or_watch"]] for r in risk_summary])

# ── 6. 加权共识结果 ──
h1("六、加权共识结果")
cons_dec = Counter()
for r in consensus_summary:
    cons_dec[r.get("decision")] += int(r.get("count") or 0)
p(f"Accepted: {cons_dec.get('accepted',0)}, Challenged: {cons_dec.get('challenged',0)}, Rejected: {cons_dec.get('rejected',0)}")
h2("6.1 决策统计")
tbl(["类型", "标签", "决策", "数量", "平均分", "平均ESS", "平均q_bar"],
    [[r["group_kind"], r["label"], r["decision"], r["count"], f"{float(r['avg_score']):.4f}", f"{float(r['avg_ess']):.4f}", f"{float(r['avg_q_bar']):.4f}"] for r in consensus_summary])

# ── 7. Baseline 对比 ──
h1("七、Baseline 对比实验")
p("Hard 数据集下四种方法的性能对比：")
tbl(["方法", "准确率", "误拦率", "漏检率", "攻击成功率"],
    [[r["method"], f"{float(r['accuracy']):.4f}", f"{float(r['false_block_rate']):.4f}", f"{float(r['missed_threat_rate']):.4f}", f"{float(r['attack_success_rate']):.4f}"] for r in baseline])
p("关键发现：zero_trust_weighted_consensus 在 Hard 数据集下准确率降至 98.33%，而非 Easy 下的 100%。验证了系统在极限攻击下的真实韧性边界。")

# ── 8. 延迟 ──
h1("八、流水线延迟")
tbl(["阶段", "总耗时(ms)", "数量", "平均(ms)"],
    [[r["stage"], r["total_ms"], r["count"], r["avg_ms"]] for r in latency])

# ── 9. Easy vs Hard 对照 ──
h1("九、Easy vs Hard 关键对比")
tbl(["维度", "Easy", "Hard"],
    [["Agent总数", "54", "90"],
     ["Task总数", "120", "300"],
     ["声明总数", "1,200", "4,890"],
     ["攻击家族数", "5", "7"],
     ["Byzantine信任度", "42-61", "72-86"],
     ["投毒证据质量", "18-48", "65-82"],
     ["假证据质量", "无", "85-98 (匹敌真证据)"],
     ["Validation通过率", "80.6%", "97.6%"],
     ["BSS区分度", "明显 (可检测)", "盲区 (全部<0.15)"],
     ["zero_trust准确率", "100%", "98.33%"],
     ["majority_vote准确率", "0%", "30.67%"],
     ["direct_trust准确率", "29.17%", "34.33%"],
     ["single_verifier准确率", "100%", "60.00%"]])

# ── 10. 结果文件清单 ──
h1("十、结果文件证据清单")
tbl(["文件", "用途"],
    [["实验数据集/mabzt_comm_dataset_hard/", "Hard 数据集（7种攻击家族）"],
     ["实验结果/hard/baseline_comparison.csv", "四方法准确率对比"],
     ["实验结果/hard/validation_summary.csv", "七项零信任校验失败率"],
     ["实验结果/hard/risk_scores.csv", "批处理 BSS 风险结果"],
     ["实验结果/hard/consensus_results.csv", "加权共识决策详情"],
     ["实验结果/hard/consensus_decision_summary.csv", "共识决策汇总"],
     ["实验结果/hard/latency_results.csv", "流水线各阶段延迟"],
     ["实验结果/hard/run_summary.json", "运行摘要"],
     ["实验结果/hard/graph_snapshot.json", "声明图谱快照"],
     ["实验结果/hard/neo4j_import.cypher", "Neo4j 导入脚本"]])

doc.save(str(OUT))
print(f"Generated: {OUT}  ({OUT.stat().st_size} bytes)")
