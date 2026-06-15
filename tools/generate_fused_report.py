#!/usr/bin/env python3
"""生成多 Agent 零信任协同与知识投毒因果验证融合版作品报告。"""

from __future__ import annotations

import os
import json
import sys
from pathlib import Path
from typing import Iterable, Sequence

ROOT = Path(__file__).resolve().parents[2]
DEPS = ROOT / "zyjd_system" / ".docx_deps"
sys.path.insert(0, str(DEPS))

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

OUT_DIR = ROOT / "zyjd_system" / "docs" / "fused_report"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "基于多Agent零信任协同与RAG知识投毒因果验证融合的信息污染溯源纠偏系统.docx"
EXPERIMENT_RESULT = ROOT / "zyjd_system" / "experiments" / "results" / "report_experiments.json"
EXPERIMENT_FIGURE_DIR = ROOT / "zyjd_system" / "experiments" / "figures"

COLORS = {
    "blue": "#DDEBF7",
    "orange": "#FCE4D6",
    "green": "#E2F0D9",
    "purple": "#E4DFEC",
    "red": "#F4CCCC",
    "gray": "#F2F2F2",
    "yellow": "#FFF2CC",
}


def configure_matplotlib() -> None:
    """配置可用的中文字体与统一图形风格。"""
    windows_font = Path("/mnt/c/Windows/Fonts/simsun.ttc")
    if windows_font.exists():
        font_manager.fontManager.addfont(str(windows_font))
        plt.rcParams["font.family"] = font_manager.FontProperties(fname=str(windows_font)).get_name()
    else:
        plt.rcParams["font.sans-serif"] = ["DejaVu Sans", "Droid Sans Fallback"]
    plt.rcParams["axes.unicode_minus"] = False


def add_box(
    ax,
    x: float,
    y: float,
    w: float,
    h: float,
    text: str,
    color: str,
    fontsize: float = 12,
    linewidth: float = 1.2,
) -> FancyBboxPatch:
    """添加圆角矩形节点。"""
    patch = FancyBboxPatch(
        (x, y),
        w,
        h,
        boxstyle="round,pad=0.015,rounding_size=0.025",
        facecolor=color,
        edgecolor="#333333",
        linewidth=linewidth,
    )
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, color="#111111")
    return patch


def connect(
    ax,
    start: tuple[float, float],
    end: tuple[float, float],
    color: str = "#333333",
    style: str = "-|>",
    width: float = 1.2,
    rad: float = 0.0,
) -> None:
    """在节点边框坐标之间绘制箭头，避免箭头进入节点内部。"""
    arrow = FancyArrowPatch(
        start,
        end,
        arrowstyle=style,
        mutation_scale=12,
        linewidth=width,
        color=color,
        connectionstyle=f"arc3,rad={rad}",
        shrinkA=0,
        shrinkB=0,
    )
    ax.add_patch(arrow)


def setup_canvas(figsize: tuple[float, float] = (13, 6.2)):
    fig, ax = plt.subplots(figsize=figsize, dpi=180)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    return fig, ax


def save_figure(fig, name: str) -> Path:
    path = ASSET_DIR / name
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def figure_background() -> Path:
    fig, ax = setup_canvas()
    add_box(ax, 0.03, 0.67, 0.18, 0.19, "多 Agent 协同\n规划、检索、验证、执行", COLORS["purple"], 13)
    add_box(ax, 0.03, 0.22, 0.18, 0.19, "协同污染入口\n幻觉、篡改、拜占庭节点", COLORS["red"], 12)
    add_box(ax, 0.31, 0.67, 0.18, 0.19, "零信任声明链\n身份、权限、签名、证据", COLORS["green"], 13)
    add_box(ax, 0.31, 0.22, 0.18, 0.19, "开放网页 / 企业文档\n知识库与工具返回", COLORS["blue"], 12)
    add_box(ax, 0.59, 0.67, 0.18, 0.19, "知识投毒因果验证\nRAS / GIS / 反事实", COLORS["orange"], 13)
    add_box(ax, 0.59, 0.22, 0.18, 0.19, "联合信息污染溯源\n证据—声明—消息—动作", COLORS["blue"], 12)
    add_box(ax, 0.82, 0.445, 0.15, 0.19, "可信纠偏与\n安全恢复", COLORS["yellow"], 13)
    connect(ax, (0.21, 0.765), (0.31, 0.765))
    connect(ax, (0.21, 0.315), (0.31, 0.315), color="#B22222")
    connect(ax, (0.40, 0.41), (0.40, 0.67), color="#B22222")
    connect(ax, (0.49, 0.765), (0.59, 0.765))
    connect(ax, (0.49, 0.315), (0.59, 0.315), color="#B22222")
    connect(ax, (0.68, 0.67), (0.68, 0.41), color="#B22222")
    connect(ax, (0.77, 0.765), (0.82, 0.59))
    connect(ax, (0.77, 0.315), (0.82, 0.49), color="#B22222")
    return save_figure(fig, "fig1_background.png")


def figure_overall_flow() -> Path:
    fig, ax = setup_canvas((14, 7.0))
    labels = [
        ("多 Agent\n任务协同", COLORS["purple"]),
        ("零信任声明包\n身份·权限·签名·证据", COLORS["green"]),
        ("知识证据接入\n网页·文档·工具结果", COLORS["blue"]),
        ("RAS / GIS\n双条件风险检测", COLORS["orange"]),
        ("四路反事实\n因果验证", COLORS["purple"]),
        ("联合污染溯源\n源头与传播路径", COLORS["blue"]),
        ("隔离·回滚\n可信共识重建", COLORS["red"]),
        ("可信重生成\n复核与风险报告", COLORS["yellow"]),
    ]
    xs = [0.04, 0.285, 0.53, 0.775]
    w, h = 0.185, 0.19
    top_y, bottom_y = 0.68, 0.28
    positions = [
        (xs[0], top_y),
        (xs[1], top_y),
        (xs[2], top_y),
        (xs[3], top_y),
        (xs[3], bottom_y),
        (xs[2], bottom_y),
        (xs[1], bottom_y),
        (xs[0], bottom_y),
    ]
    for (label, color), (x, y) in zip(labels, positions):
        add_box(ax, x, y, w, h, label, color, 13)
    for i in range(3):
        x, y = positions[i]
        nx, ny = positions[i + 1]
        connect(ax, (x + w, y + h / 2), (nx, ny + h / 2))
    connect(ax, (xs[3] + w / 2, top_y), (xs[3] + w / 2, bottom_y + h))
    for i in range(4, 7):
        x, y = positions[i]
        nx, ny = positions[i + 1]
        connect(ax, (x, y + h / 2), (nx + w, ny + h / 2))
    feedback_y = 0.12
    outside_x = 0.985
    connect(ax, (xs[0] + w / 2, bottom_y), (xs[0] + w / 2, feedback_y), color="#B22222", style="-")
    connect(ax, (xs[0] + w / 2, feedback_y), (outside_x, feedback_y), color="#B22222", style="-")
    connect(ax, (outside_x, feedback_y), (outside_x, top_y + h / 2), color="#B22222", style="-")
    connect(ax, (outside_x, top_y + h / 2), (xs[3] + w, top_y + h / 2), color="#B22222")
    ax.text(0.51, 0.045, "复核未通过", ha="center", fontsize=13, color="#8B0000")
    return save_figure(fig, "fig2_overall_flow.png")


def figure_framework() -> Path:
    fig, ax = setup_canvas((13.5, 6.7))
    layers = [
        ("交互与展示层", ["案例运行", "传播图谱", "风险报告", "纠偏前后对比"], COLORS["blue"]),
        ("可信处置层", ["证据加权共识", "Chunk / Agent 隔离", "状态回滚", "可信重生成"], COLORS["green"]),
        ("检测与因果层", ["RAS / GIS / DualRisk", "四路反事实", "CausalScore", "声明—证据 NLI"], COLORS["orange"]),
        ("零信任与溯源层", ["ZT-Claim Envelope", "身份与权限校验", "证据哈希", "联合 Provenance Graph"], COLORS["purple"]),
        ("数据与编排层", ["网页 / 企业知识库", "检索器", "多 Agent 编排", "工具调用与 Checkpoint"], COLORS["gray"]),
    ]
    y_values = [0.80, 0.62, 0.44, 0.26, 0.08]
    for (layer, modules, color), y in zip(layers, y_values):
        add_box(ax, 0.03, y, 0.15, 0.12, layer, color, 12.5)
        module_w = 0.175
        for i, module in enumerate(modules):
            x = 0.22 + i * 0.19
            add_box(ax, x, y, module_w, 0.12, module, "white", 11.2, 1.0)
            if i < len(modules) - 1:
                connect(ax, (x + module_w, y + 0.06), (x + 0.19, y + 0.06), color="#777777", style="-")
        if y != y_values[-1]:
            connect(ax, (0.50, y), (0.50, y - 0.05), color="#555555")
    return save_figure(fig, "fig3_framework.png")


def figure_provenance() -> Path:
    fig, ax = setup_canvas((14, 6.3))
    nodes = {
        "page": (0.04, 0.70, 0.12, 0.12, "Page / Document", COLORS["blue"]),
        "chunk": (0.23, 0.70, 0.12, 0.12, "Chunk", COLORS["orange"]),
        "query": (0.23, 0.30, 0.12, 0.12, "Query", COLORS["gray"]),
        "evidence": (0.42, 0.70, 0.12, 0.12, "Evidence", COLORS["green"]),
        "claim": (0.60, 0.70, 0.12, 0.12, "Claim", COLORS["yellow"]),
        "agent1": (0.42, 0.30, 0.12, 0.12, "Agent A", COLORS["purple"]),
        "agent2": (0.60, 0.30, 0.12, 0.12, "Agent B", COLORS["red"]),
        "answer": (0.80, 0.70, 0.12, 0.12, "Answer / Action", COLORS["green"]),
        "checkpoint": (0.80, 0.30, 0.12, 0.12, "Checkpoint", COLORS["blue"]),
    }
    for x, y, w, h, text, color in nodes.values():
        add_box(ax, x, y, w, h, text, color, 12)
    connect(ax, (0.16, 0.76), (0.23, 0.76))
    connect(ax, (0.29, 0.42), (0.29, 0.70))
    connect(ax, (0.35, 0.76), (0.42, 0.76))
    connect(ax, (0.54, 0.76), (0.60, 0.76))
    connect(ax, (0.72, 0.76), (0.80, 0.76))
    connect(ax, (0.48, 0.42), (0.48, 0.70))
    connect(ax, (0.54, 0.36), (0.60, 0.36), color="#B22222")
    connect(ax, (0.66, 0.42), (0.66, 0.70), color="#B22222")
    connect(ax, (0.86, 0.70), (0.86, 0.42))
    rollback_y = 0.18
    connect(ax, (0.80, 0.36), (0.76, 0.36), color="#B22222", style="-")
    connect(ax, (0.76, 0.36), (0.76, rollback_y), color="#B22222", style="-")
    connect(ax, (0.76, rollback_y), (0.66, rollback_y), color="#B22222", style="-")
    connect(ax, (0.66, rollback_y), (0.66, 0.30), color="#B22222")
    ax.text(0.19, 0.78, "contains", fontsize=10)
    ax.text(0.30, 0.55, "retrieved_by", fontsize=10)
    ax.text(0.37, 0.78, "bound_to", fontsize=10)
    ax.text(0.55, 0.78, "supports / contradicts", fontsize=10)
    ax.text(0.74, 0.78, "caused_error", fontsize=10, color="#8B0000")
    ax.text(0.49, 0.55, "produces", fontsize=10)
    ax.text(0.55, 0.38, "message / derives", fontsize=10, color="#8B0000")
    ax.text(0.67, 0.55, "amplifies", fontsize=10, color="#8B0000")
    ax.text(0.87, 0.55, "snapshot", fontsize=10)
    ax.text(0.71, 0.14, "rollback / replan", fontsize=10, color="#8B0000")
    return save_figure(fig, "fig4_provenance.png")


def figure_correction() -> Path:
    fig, ax = setup_canvas((13.5, 6.2))
    items = [
        (0.08, 0.62, "风险确认\nDualRisk + CausalScore", COLORS["orange"]),
        (0.36, 0.72, "隔离与降权\nChunk / Claim / Agent", COLORS["red"]),
        (0.65, 0.62, "回滚与重规划\n恢复可信 Checkpoint", COLORS["purple"]),
        (0.72, 0.27, "风险感知重检索\n独立低风险证据", COLORS["blue"]),
        (0.39, 0.17, "可信重生成\n声明级引用", COLORS["green"]),
        (0.08, 0.27, "复核与审计\nNLI / TrustScore / 报告", COLORS["yellow"]),
    ]
    boxes = []
    for x, y, text, color in items:
        boxes.append((x, y, 0.20, 0.14))
        add_box(ax, x, y, 0.20, 0.14, text, color, 12.2)
    connect(ax, (0.28, 0.69), (0.36, 0.79))
    connect(ax, (0.56, 0.79), (0.65, 0.69))
    connect(ax, (0.75, 0.62), (0.82, 0.41))
    connect(ax, (0.72, 0.34), (0.59, 0.24))
    connect(ax, (0.39, 0.24), (0.28, 0.34))
    connect(ax, (0.18, 0.41), (0.18, 0.62), color="#B22222")
    ax.text(0.18, 0.50, "未通过则继续处置", ha="center", fontsize=11, color="#8B0000")
    return save_figure(fig, "fig5_correction.png")


def figure_consensusguard() -> Path:
    """绘制 ConsensusGuard 独立板块流程图。"""
    fig, ax = setup_canvas((14, 5.8))
    labels = [
        ("多 Agent 编排\n规划·分析·验证·执行", COLORS["purple"]),
        ("零信任声明包\n身份·权限·证据·签名", COLORS["green"]),
        ("声明传播图\nClaim Provenance DAG", COLORS["blue"]),
        ("级联错误检测\n传播·漂移·拜占庭嫌疑", COLORS["orange"]),
        ("可信共识与纠偏\n隔离·回滚·重规划", COLORS["yellow"]),
    ]
    x_values = [0.025, 0.225, 0.425, 0.625, 0.825]
    y, w, h = 0.48, 0.15, 0.24
    for (label, color), x in zip(labels, x_values):
        add_box(ax, x, y, w, h, label, color, 12)
    for i in range(len(x_values) - 1):
        connect(ax, (x_values[i] + w, y + h / 2), (x_values[i + 1], y + h / 2))
    feedback_y = 0.22
    connect(ax, (x_values[-1] + w / 2, y), (x_values[-1] + w / 2, feedback_y), color="#B22222", style="-")
    connect(ax, (x_values[-1] + w / 2, feedback_y), (x_values[0] + w / 2, feedback_y), color="#B22222", style="-")
    connect(ax, (x_values[0] + w / 2, feedback_y), (x_values[0] + w / 2, y), color="#B22222")
    ax.text(0.5, 0.14, "纠偏后重新进入可信协同", ha="center", fontsize=12, color="#8B0000")
    return save_figure(fig, "fig_consensusguard.png")


def figure_zhiyuan_rag() -> Path:
    """绘制智源净域 RAG 知识投毒防御流程图。"""
    fig, ax = setup_canvas((14, 5.8))
    labels = [
        ("RAG 知识源\n网页·企业文档·工具结果", COLORS["blue"]),
        ("检索与证据对象\nChunk·排名·来源·哈希", COLORS["green"]),
        ("双条件检测\nRAS·GIS·DualRisk", COLORS["orange"]),
        ("四路反事实验证\nCausalScore", COLORS["purple"]),
        ("隔离与可信重生成\n重检索·NLI·TrustScore", COLORS["yellow"]),
    ]
    x_values = [0.025, 0.225, 0.425, 0.625, 0.825]
    y, w, h = 0.48, 0.15, 0.24
    for (label, color), x in zip(labels, x_values):
        add_box(ax, x, y, w, h, label, color, 12)
    for i in range(len(x_values) - 1):
        connect(ax, (x_values[i] + w, y + h / 2), (x_values[i + 1], y + h / 2))
    feedback_y = 0.22
    connect(ax, (x_values[-1] + w / 2, y), (x_values[-1] + w / 2, feedback_y), color="#B22222", style="-")
    connect(ax, (x_values[-1] + w / 2, feedback_y), (x_values[2] + w / 2, feedback_y), color="#B22222", style="-")
    connect(ax, (x_values[2] + w / 2, feedback_y), (x_values[2] + w / 2, y), color="#B22222")
    ax.text(0.70, 0.14, "重生成复核未通过", ha="center", fontsize=12, color="#8B0000")
    return save_figure(fig, "fig_zhiyuan_rag.png")


def set_run_font(run, east_asia: str = "宋体", size: float = 12, bold: bool = False) -> None:
    run.font.name = "SimSun"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.5)

    for style in doc.styles:
        try:
            style.font.name = "SimSun"
            style._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
            style.font.color.rgb = RGBColor(0, 0, 0)
        except (AttributeError, ValueError):
            continue

    normal = doc.styles["Normal"]
    normal.font.name = "SimSun"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for name, size in (
        ("Title", 22),
        ("Heading 1", 16),
        ("Heading 2", 14),
        ("Heading 3", 12),
        ("Heading 4", 11),
    ):
        style = doc.styles[name]
        style.font.name = "SimSun"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Caption"].font.size = Pt(10.5)
    doc.styles["Caption"].font.color.rgb = RGBColor(0, 0, 0)
    doc.styles["Caption"]._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    doc.styles["Caption"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.styles["Caption"].paragraph_format.first_line_indent = Pt(0)

    settings = doc.settings._element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, size=10.5)
    for field_type, text in (("begin", None), (None, " PAGE "), ("end", None)):
        if field_type:
            node = OxmlElement("w:fldChar")
            node.set(qn("w:fldCharType"), field_type)
        else:
            node = OxmlElement("w:instrText")
            node.set(qn("xml:space"), "preserve")
            node.text = text
        run._r.append(node)


def restart_page_number(section, start: int = 1) -> None:
    """在正文分节中重新开始页码。"""
    section.footer.is_linked_to_previous = False
    footer = section.footer.paragraphs[0]
    footer.clear()
    add_page_number(footer)
    pg_num_type = section._sectPr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        section._sectPr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-4" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "目录将在 Word/WPS 中更新后显示"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, placeholder, end])


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))


def add_list(doc: Document, items: Iterable[str], numbered: bool = False) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph(style="Normal" if numbered else "List Bullet")
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(24)
        text = f"{index}. {item}" if numbered else item
        set_run_font(p.add_run(text))


def set_cell_border(cell, top: int | None = None, bottom: int | None = None) -> None:
    """设置三线表单元格边框：仅允许顶线、表头线和底线。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        size = top if edge == "top" else bottom if edge == "bottom" else None
        element.set(qn("w:val"), "single" if size else "nil")
        if size:
            element.set(qn("w:sz"), str(size))
        element.set(qn("w:color"), "000000")


def add_table(
    doc: Document,
    number: int,
    title: str,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
) -> None:
    caption = doc.add_paragraph(f"表 {number}  {title}", style="Caption")
    caption.paragraph_format.keep_with_next = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_run_font(p.add_run(value), size=10.5, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_border(cell, top=12, bottom=6)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            set_run_font(p.add_run(str(value)), size=10.5)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cells[i])
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom=12)
    doc.add_paragraph()


def add_figure(doc: Document, path: Path, number: int, caption: str, width_cm: float = 15.5) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph(f"图 {number}  {caption}", style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_formula(doc: Document, formula: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run(formula)
    run.font.name = "SimSun"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("第十九届全国大学生信息安全竞赛（作品赛）"), "宋体", 18, True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("暨第三届“长城杯”网数智安全大赛（作品赛）"), "宋体", 16, True)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("作 品 报 告"), "宋体", 26, True)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("□命题赛道             ■自由赛道"), "宋体", 12)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("作品名称："), "宋体", 14, True)
    title = "基于多 Agent 零信任协同与 RAG 知识投毒因果验证融合的\n信息污染溯源纠偏系统"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run(title), "宋体", 20, True)
    for _ in range(5):
        doc.add_paragraph()
    for label in ("电子邮箱：____________________________", "提交日期：____________________________"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_run_font(p.add_run(label), "宋体", 12)
    doc.add_page_break()


def build_report(figures: Sequence[Path]) -> None:
    doc = Document()
    setup_document(doc)
    add_page_number(doc.sections[0].footer.paragraphs[0])
    add_cover(doc)

    doc.add_heading("摘  要", level=1)
    add_body(
        doc,
        "多 Agent 系统正在把规划、检索、分析、验证和执行分配给不同智能体，并通过消息共享共同形成决策。"
        "如果 Agent 的身份、权限、声明和证据被默认信任，单个节点的幻觉、通信篡改或拜占庭行为就可能沿协作链"
        "扩散，形成错误共识和高风险动作。与此同时，AI 搜索与企业检索增强生成（RAG）持续引入外部知识和工具"
        "结果，定向投毒片段可能先通过检索吸附进入上下文，再诱导 Agent 生成错误声明。现有方案通常只"
        "覆盖知识侧过滤、最终答案核验或 Agent 侧投票，难以统一回答“污染从哪里进入、如何跨证据与消息传播、"
        "哪些声明真正导致错误、如何恢复可信状态”等问题。"
    )
    add_body(
        doc,
        "系统首先以零信任声明包约束 Agent 协同，为每条关键 Claim 绑定身份、角色权限、父子依赖、证据哈希和"
        "签名信息；"
        "随后将网页、企业文档、检索 Chunk 和工具输出统一映射为 Evidence，使用 RAS、GIS 与 DualRisk 发现候选"
        "投毒片段，并通过原始、删除可疑、仅可疑、可信替代四路反事实生成计算 CausalScore。"
    )
    add_body(
        doc,
        "在完成零信任校验和因果验证后，系统构建 Page—Document—Chunk—Query—Evidence—Claim—Agent—"
        "Answer/Action 联合溯源图，"
        "同时追踪知识复制、声明支持与矛盾、消息派生、Agent 放大和错误动作依赖；通过声明—证据矩阵、来源独立性、"
        "历史信誉、角色权限与因果风险形成证据加权可信共识。对高风险 Chunk、Claim 或 Agent 执行降权、隔离、"
        "回滚和重规划，随后进行风险感知重检索、可信重生成与二次复核，形成“零信任接入—双条件检测—因果确认—"
        "联合溯源—可信共识—隔离回滚—重生成”的完整闭环。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(
        p.add_run(
            "关键词：AI 搜索；检索增强生成；知识投毒；多智能体系统；零信任；级联错误；反事实验证；可信纠偏"
        ),
        bold=True,
    )
    doc.add_page_break()
    doc.add_heading("目  录", level=1)
    add_toc(doc)
    doc.add_page_break()

    doc.add_heading("第一章 作品概述", level=1)
    doc.add_heading("1.1 背景分析", level=2)
    add_body(
        doc,
        "大模型应用正由单体问答演进为多 Agent 协同。不同 Agent 负责规划、检索、分析、验证和执行，其输出会被"
        "其他 Agent 引用并写入共享状态。多智能体系统的失败研究揭示了角色错配、协作失配、验证和终止机制不足等"
        "系统性问题[31]；通信攻击与 Prompt Infection 进一步说明，攻击者可以利用消息链路和上下文共享使污染在"
        "多个 Agent 之间传播[32-33]。因此，协同过程首先需要建立“持续验证、最小权限、声明有据”的零信任机制。"
    )
    add_body(
        doc,
        "在零信任协同的基础上，还必须验证 Agent 使用的知识。AI 搜索与企业 RAG 通过检索外部网页或内部文档生成"
        "带引用答案，但也使共享目录、上传文档和第三方资料成为新的知识入口。PoisonedRAG 表明，有效知识投毒通常"
        "同时满足“进入 Top-K”的检索条件和“改变最终答案”的生成条件，少量定向文本即可产生显著影响[6]。"
    )
    add_figure(doc, figures[0], 1, "复合信息污染背景与风险扩散路径")
    add_body(
        doc,
        "知识投毒与多 Agent 级联错误并非两条孤立风险链。被污染的检索片段可能成为 Agent 的证据输入，错误 Agent "
        "也可能反复引用同一污染源并影响其他节点。因此，本作品把知识入口、声明传播和协作状态纳入同一溯源模型，"
        "以声明级因果关系而非单纯文本异常或多数票作为可信判断依据。"
    )

    doc.add_heading("1.2 相关工作", level=2)
    add_body(
        doc,
        "在多 Agent 安全方面，Why Do Multi-Agent LLM Systems Fail? 建立了多智能体失败模式分类[31]；"
        "Agent-in-the-Middle 与 Prompt Infection 说明通信链和外部内容可成为系统级攻击面[32-33]；G-Safeguard "
        "采用拓扑视角分析多 Agent 风险[34]。NIST SP 800-207 强调持续验证而非默认信任[30]，为 Agent 间声明校验"
        "提供了安全原则。"
    )
    add_body(
        doc,
        "在知识投毒方面，PoisonedRAG、Pandora、Phantom 与 Poisoned-MRAG 分别研究定向知识污染、间接越权、"
        "触发式后门和多模态投毒[6-9]。RAGForensics、RobustRAG 及相关安全 RAG 研究覆盖候选溯源、隔离聚合和"
        "防御评测[10-14]。ALCE、RAGAS、ARES、FActScore 与 FEVER 为引用支持、忠实性和声明验证提供了方法基础"
        "[17-21]，但质量评价仍需结合投毒风险与来源独立性。"
    )
    add_body(
        doc,
        "在 RAG 基础研究方面，Lewis 等提出参数化生成与非参数化检索结合的 RAG[1]，DPR、Natural Questions、"
        "HotpotQA 和 MS MARCO 奠定了检索与开放域问答评测基础[2-5]。GEO 研究内容表达对生成式搜索可见性的"
        "影响[15]，同时也揭示伪权威表达和站群互引可能被滥用。"
    )
    add_table(
        doc,
        1,
        "两类研究与本作品融合改进",
        ["研究方向", "已有能力", "主要不足", "本作品融合点"],
        [
            ["多 Agent 安全", "失败分类、通信攻击、拓扑检测", "常默认输入知识本身可信", "将 Chunk 风险和证据因果关系纳入声明链"],
            ["RAG 投毒防御", "过滤、鲁棒聚合、答案溯源", "难覆盖 Agent 间传播与状态回滚", "因果确认结果进入协同图与纠偏引擎"],
            ["事实与引用评测", "声明拆分、支持/矛盾判断", "忠实复述投毒上下文仍可能得高分", "TrustScore 联合来源独立性与投毒风险"],
            ["零信任架构", "持续身份与访问验证", "未直接定义 LLM 声明可信度", "扩展为带证据、依赖和签名的声明包"],
        ],
    )

    doc.add_heading("1.3 作品目标与应用前景", level=2)
    add_list(
        doc,
        [
            "统一发现知识侧投毒和协同侧级联错误，避免只在最终答案阶段被动核验。",
            "定位污染源 Chunk、错误 Claim、异常通信边和可疑 Agent，解释传播路径与实际影响。",
            "以独立证据和反事实结果重建可信共识，并支持隔离、回滚、重规划和可信重生成。",
            "面向企业知识库问答、AI 搜索可信引用、安全运营协同和高风险决策辅助提供可审计防护。",
        ],
        numbered=True,
    )

    doc.add_heading("第二章 作品设计与实现", level=1)
    doc.add_heading("2.1 多 Agent 零信任协同", level=2)
    add_body(
        doc,
        "系统首先解决多 Agent 协同中的默认信任问题。规划、检索、分析、验证和执行 Agent 之间不直接传递不可审计"
        "的自由文本结论，而是使用零信任声明包传递原子 Claim。每条声明必须绑定 Agent 身份、角色权限、父声明、"
        "证据引用、工具调用哈希、nonce、时间戳和签名；校验不通过的消息不得进入共享状态和后续共识。"
    )
    add_table(
        doc,
        2,
        "多 Agent 零信任协同对象",
        ["对象", "关键字段", "零信任作用"],
        [
            ["Agent Identity", "agent_id、role、permission、public_key", "确认节点身份并实施最小权限"],
            ["Claim", "subject、predicate、object、confidence、parent_claims", "将自然语言结论拆为可验证声明"],
            ["Evidence", "source、chunk_id、content_hash、tool_hash、scores", "要求声明绑定可追踪证据"],
            ["ZT-Claim Envelope", "claim、evidence_refs、nonce、timestamp、signature", "校验完整性、防重放和越权"],
            ["Checkpoint", "task_id、round、trusted_state_hash、dependencies", "为污染状态回滚提供可信锚点"],
        ],
    )
    add_body(
        doc,
        "系统不把 Agent 自报置信度直接当作可信度。通过身份与签名校验后，还需检查角色是否有权生成该类声明、"
        "证据是否存在、父子依赖是否闭合、工具结果是否与哈希一致。对无证据高置信声明、通信篡改、角色越权和持续"
        "误导行为分别执行拒绝、挑战补证、降权或隔离。"
    )
    add_body(
        doc,
        "在可信共识阶段，Agent 权重由证据支持度、历史信誉、角色权限、独立验证结果和传播路径共同决定，避免多个"
        "Agent 重复引用同一错误源后形成虚假多数。高风险执行动作默认设置人工确认门。"
    )

    doc.add_heading("2.2 知识投毒检测与因果验证", level=2)
    add_body(
        doc,
        "零信任协同只能保证消息身份与传输过程可验证，不能保证 Agent 引用的知识本身正确。因此，系统第二步对"
        "网页、企业文档、检索 Chunk 和工具结果进行知识投毒检测。所有知识对象统一为 Evidence，并记录来源、"
        "内容哈希、检索排名、引用状态和风险评分。"
    )
    add_body(doc, "检索吸附性分数衡量某 Chunk 被检索的频率相对随机基线的放大程度：")
    add_formula(doc, "RAS(chunkᵢ) = (freqᵢ / total_retrievals) / (1 / total_chunks)")
    add_body(
        doc,
        "其中 freqᵢ 为 Chunk 在历史检索中出现次数；RAS 大于 1 表示其出现频率高于随机基线。答案诱导性分数衡量"
        "答案对候选 Chunk 的相对依赖："
    )
    add_formula(doc, "GIS(chunkᵢ, answer) = sim(answer, chunkᵢ) / max sim(answer, all_chunks)")
    add_body(doc, "MVP 使用 TF-IDF 余弦相似度，并以乘积形成双条件风险：")
    add_formula(doc, "DualRisk(chunkᵢ) = NormalizedRAS(chunkᵢ) × GIS(chunkᵢ)")
    add_body(
        doc,
        "高 DualRisk 只表示候选可疑，不能直接证明致错。系统固定查询、生成模板和解码设置，比较原始 Top-K、"
        "删除可疑、仅使用可疑、可信证据替代四路答案。若原始答案接近仅可疑答案，而删除或替代后恢复正确结论，"
        "则候选具有较强因果贡献。"
    )
    add_formula(doc, "CausalScore = 1 - sim(A_orig, A_remove) / max(sim(A_orig, A_only), ε)")

    doc.add_heading("2.3 信息污染联合溯源", level=2)
    add_figure(doc, figures[3], 2, "知识投毒与多 Agent 级联错误联合溯源图")
    add_body(
        doc,
        "系统第三步把零信任声明校验结果和知识投毒因果验证结果写入同一溯源图。"
        "联合图以 Page、Document、Chunk、Query、Evidence、Claim、Agent、Answer/Action 和 Checkpoint 为节点，"
        "以 contains、retrieved_by、supports、contradicts、copied_from、same_claim、derives、amplifies、"
        "caused_error、isolated_in 和 rollback_to 等关系记录信息流。图谱既能识别多个网页是否复制同一错误声明，"
        "也能判断哪些 Agent 在缺乏独立证据的情况下放大了该声明。"
    )
    add_table(
        doc,
        3,
        "联合图关键关系及安全语义",
        ["关系", "方向", "安全语义"],
        [
            ["retrieved_by", "Chunk → Query", "记录可疑内容进入上下文的检索链"],
            ["supports / contradicts", "Evidence → Claim", "记录证据对声明的支持或矛盾"],
            ["copied_from / same_claim", "Document/Claim ↔ Document/Claim", "识别站群、转载与虚假多源一致性"],
            ["derives / amplifies", "Claim/Agent → Claim/Agent", "追踪声明派生和级联放大"],
            ["caused_error", "Chunk/Claim → Answer/Action", "记录经反事实确认的致错关系"],
            ["rollback_to", "Task → Checkpoint", "记录可信状态恢复位置"],
        ],
    )

    doc.add_heading("2.4 可信纠偏与安全恢复", level=2)
    add_body(
        doc,
        "系统第四步依据溯源路径确定纠偏范围。对知识侧高风险 Chunk 执行隔离或降权，对无证据 Claim 触发挑战补证，"
        "对篡改消息直接拒绝，对持续误导或越权的 Agent 限制其参与高风险决策。若污染已进入共享状态，则沿依赖图"
        "计算受影响节点并回滚到最近可信 Checkpoint，再将任务重新分配给来源独立的 Agent。"
    )
    add_figure(doc, figures[4], 3, "可信纠偏与证据化恢复闭环")
    add_body(
        doc,
        "回滚后扩大检索候选池，过滤已隔离内容，并加入来源多样性和风险惩罚重新排序。可信重生成只使用通过声明级"
        "裁决的证据，为每个关键 Claim 绑定引用；新答案再次接受 NLI、TrustScore 和图谱一致性复核，未通过则返回"
        "知识投毒检测与因果验证阶段。"
    )

    doc.add_heading("2.5 声明级可信评分", level=2)
    add_body(
        doc,
        "系统把候选答案拆为原子 Claim，使用 NLI/启发式规则判断每条 Evidence 与 Claim 的支持、矛盾或无关关系，"
        "形成声明—证据矩阵。来源独立性通过域名、文档哈希、copied_from 与 same_claim 关系估计。"
    )
    add_formula(
        doc,
        "TrustScore = 0.30×EvidenceSupportRate + 0.20×SourceIndependence "
        "+ 0.30×(1-NormalizedDualRisk) + 0.20×(1-NormalizedCausalRisk)",
    )
    add_body(
        doc,
        "该评分将证据支持、来源独立性、知识投毒风险和反事实因果风险统一到 [0,1]，作为可信共识、处置级别和"
        "重生成结果复核的共同依据。"
    )

    doc.add_heading("2.6 系统总体流程", level=2)
    add_figure(doc, figures[1], 4, "系统总体流程")
    add_body(
        doc,
        "系统先建立多 Agent 任务协同与零信任声明包，再接入知识证据并完成双条件检测和四路反事实因果验证，"
        "随后构建联合污染溯源图，最后执行隔离、回滚、可信共识重建和可信重生成。"
    )

    doc.add_heading("2.7 系统实现框架", level=2)
    add_figure(doc, figures[2], 5, "五层协同防御实现框架")
    add_body(
        doc,
        "MVP 采用 Python、FastAPI、Pydantic、scikit-learn 与 networkx 实现。文本相似度使用 TF-IDF 余弦相似度，"
        "声明—证据关系采用启发式 NLI，联合图采用 networkx 异构属性图；重生成使用受约束模板，不依赖 GPU、"
        "在线 API 或长时间训练。各服务通过 Protocol/ABC 暴露接口，可在后续替换底层实现。"
    )

    doc.add_heading("2.8 后端接口与模块划分", level=2)
    add_table(
        doc,
        4,
        "主要服务模块",
        ["模块组", "核心服务", "输出"],
        [
            ["协同安全", "ClaimEnvelope、IdentityVerifier、ByzantineScorer", "声明校验与 Agent 嫌疑"],
            ["知识侧检测", "ChunkManager、RAS、GIS、DualRisk", "候选风险 Chunk 与评分"],
            ["因果确认", "Counterfactual、CausalScore", "四路答案与因果贡献"],
            ["联合溯源", "PoisonGraph、ClaimProvenanceGraph", "污染源、传播路径与影响范围"],
            ["可信裁决", "NLIJudge、ClaimEvidenceMatrix、TrustScore、ConsensusEngine", "可信 Claim 集与共识"],
            ["恢复处置", "Isolation、Rollback、RiskAwareRetrieval、TrustedRegeneration", "恢复状态、可信答案与报告"],
        ],
    )
    add_body(
        doc,
        "FastAPI 对外提供健康检查、案例列表、完整流水线运行、检索、检测、反事实、图谱、可信评分和可信重生成接口。"
        "所有接口使用统一 success/data/error/timestamp 包装，便于前端大屏展示风险链和处置结果。"
    )

    doc.add_heading("第三章 作品测试与分析方案", level=1)
    doc.add_heading("3.1 测试原则与数据", level=2)
    add_body(
        doc,
        "本报告不填写尚未执行的实验数值。原型完成后应在固定数据版本、随机种子、检索器、模型或模板参数下运行，"
        "保存原始 JSON/CSV 结果，并将均值、标准差和置信区间回填。知识侧可组合 Natural Questions、HotpotQA、"
        "MS MARCO、FEVER、ALCE 与本地 ZhiYuan-PoisonBench 模拟案例；多 Agent 侧构造幻觉 IOC、通信篡改、"
        "恶意节点、错误多数和污染工具结果等受控场景。"
    )
    add_table(
        doc,
        5,
        "公共数据集及官方链接",
        ["数据集", "适用方向", "用途", "官方链接"],
        [
            [
                "MAST-Data",
                "多 Agent 失效与溯源",
                "1642 条多智能体执行轨迹及失效模式标注",
                "https://arxiv.org/abs/2503.13657",
            ],
            [
                "MultiAgentBench",
                "多 Agent 协同",
                "评测协作拓扑、任务完成和里程碑达成",
                "https://github.com/MultiagentBench/MARBLE",
            ],
            [
                "AgentDojo",
                "零信任工具输入",
                "评测 Agent 对不可信工具数据和间接注入的鲁棒性",
                "https://github.com/ethz-spylab/agentdojo",
            ],
            [
                "PoisonedRAG",
                "知识投毒",
                "构造受控知识污染场景并验证检测与因果确认",
                "https://arxiv.org/abs/2402.07867",
            ],
            [
                "RAGTruth",
                "RAG 幻觉检测",
                "利用响应级和词级标注评测声明忠实性",
                "https://arxiv.org/abs/2401.00396",
            ],
            [
                "ALCE",
                "AI 搜索引用可信",
                "评测答案正确性、引用完整性与引用质量",
                "https://github.com/princeton-nlp/ALCE",
            ],
            [
                "CFEVER",
                "中文事实验证",
                "中文支持、反驳和证据不足三分类",
                "https://ikmlab.github.io/CFEVER/",
            ],
        ],
    )
    add_body(
        doc,
        "上述数据集不能直接等同于本作品的完整复合威胁集。实施时应保留其原始许可证和划分，在此基础上增加本地、"
        "受控、明确标注的知识投毒与异常 Agent 注入字段，并记录转换脚本、随机种子和样本来源。"
    )
    add_table(
        doc,
        6,
        "本地模拟评测场景",
        ["场景", "污染入口", "期望验证能力"],
        [
            ["企业制度投毒", "篡改权限或审批流程文档", "Chunk 检测、因果确认、可信替代"],
            ["漏洞状态投毒", "伪造漏洞已修复或不可利用", "声明矛盾检测与来源溯源"],
            ["安全认证投毒", "伪造合规与认证结论", "来源独立性与证据支持判断"],
            ["站群投毒", "多个伪网页复制同一错误声明", "copied_from / same_claim 图谱识别"],
            ["多 Agent 级联错误", "幻觉或恶意 Agent 影响规划与执行", "异常节点定位、隔离、回滚和重规划"],
            ["良性错误负例", "过时信息、正常分歧或无意转载", "控制误报，避免把普通错误等同攻击"],
        ],
    )

    doc.add_heading("3.2 对比基线与评价指标", level=2)
    add_body(
        doc,
        "建议比较无防护系统、静态文本过滤、仅多数投票、单 Verifier Agent、仅图异常检测、仅知识投毒防御和完整"
        "融合系统。评价既覆盖知识侧检测，也覆盖协同侧恢复。"
    )
    add_table(
        doc,
        7,
        "核心评价指标",
        ["维度", "指标", "含义"],
        [
            ["检测", "Poison Precision / Recall / F1", "投毒 Chunk 识别质量"],
            ["因果", "Causal Precision / Recall / F1", "致错证据确认质量"],
            ["溯源", "Source Localization Rate", "错误源 Chunk、Claim 或 Agent 定位率"],
            ["协同", "False Consensus Rate", "错误 Claim 进入最终共识的比例"],
            ["恢复", "Correction Success Rate", "发现错误后恢复正确任务路径的比例"],
            ["安全", "High-risk Action Error Rate", "错误封禁、删除或命令执行比例"],
            ["可信", "Evidence Support Rate / TrustScore", "声明支持覆盖与综合可信度"],
            ["开销", "Latency / Token / Tool Calls", "防御引入的响应与调用成本"],
        ],
    )

    doc.add_heading("3.3 消融与边界测试", level=2)
    add_list(
        doc,
        [
            "移除 RAS 或 GIS，观察只检测单一攻击条件时的漏报变化。",
            "移除四路反事实，观察正常高相关 Chunk 被误隔离的比例。",
            "移除来源独立性，观察站群复制和错误多数对共识的影响。",
            "移除零信任签名或权限校验，验证通信篡改和角色越权检测能力。",
            "移除回滚与重规划，仅执行告警或隔离，比较纠偏成功率与任务完成率。",
            "测试空证据、单 Chunk、全投毒、全干净、正常分歧和历史信息过期等边界情况。",
        ],
    )

    doc.add_heading("第四章 创新性说明", level=1)
    add_list(
        doc,
        [
            "双入口统一威胁模型：把外部知识投毒与多 Agent 内部级联错误放入同一“证据—声明—消息—动作”链路分析。",
            "零信任声明链：将身份、权限、证据哈希、父子依赖和签名绑定到原子 Claim，避免自然语言消息被默认信任。",
            "知识投毒因果验证：以 RAS/GIS 双条件初筛和四路反事实确认区分“相关”与“真正致错”。",
            "联合传播溯源图：同时定位污染源文档、Chunk、Claim、异常通信边、可疑 Agent 和受影响动作。",
            "证据加权可信共识：突破多数投票，联合来源独立性、角色权限、历史信誉和因果风险形成可信结论。",
            "检测到恢复的闭环：支持 Chunk/Claim/Agent 隔离、状态回滚、任务重规划、风险感知重检索和可信重生成。",
        ],
        numbered=True,
    )

    doc.add_heading("第五章 总结", level=1)
    add_body(
        doc,
        "本作品融合了“智源净域”面向 AI 搜索与 RAG 的知识投毒检测、因果溯源和可信重生成能力，以及 "
        "ConsensusGuard 面向多智能体系统的零信任声明链、级联错误检测、拜占庭节点识别和回滚纠偏能力。融合后的"
        "系统不再把知识库安全和 Agent 协同安全割裂处理，而是以 Evidence、Claim、传播依赖和反事实因果关系为共同"
        "基础，形成可检测、可解释、可阻断、可恢复的完整防御闭环。"
    )
    add_body(
        doc,
        "MVP 坚持本地模拟、轻量实现和结果诚实原则，优先完成可运行、可展示、可审计的比赛原型。后续可在不改变"
        "服务接口的前提下接入更强的向量模型、NLI 模型、图学习算法和企业身份系统，并扩展到多模态知识库与更复杂"
        "的多 Agent 工作流。"
    )

    doc.add_heading("参考文献", level=1)
    references = [
        "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS. 2020.",
        "[2] KARPUKHIN V, OGUZ B, MIN S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//EMNLP. 2020.",
        "[3] KWIATKOWSKI T, PALOMAKI J, REDFIELD O, et al. Natural Questions: A Benchmark for Question Answering Research[J]. TACL, 2019, 7: 452-466.",
        "[4] YANG Z, QI P, ZHANG S, et al. HotpotQA: A Dataset for Diverse, Explainable Multi-hop Question Answering[C]//EMNLP. 2018.",
        "[5] NGUYEN T, ROSENBERG M, SONG X, et al. MS MARCO: A Human Generated Machine Reading Comprehension Dataset[C]//CoCo@NIPS. 2016.",
        "[6] ZOU W, GENG R, WANG B, JIA J. PoisonedRAG: Knowledge Corruption Attacks to Retrieval-Augmented Generation of Large Language Models[C]//USENIX Security Symposium. 2025.",
        "[7] DENG G, LIU Y, WANG K, et al. Pandora: Jailbreak GPTs by Retrieval Augmented Generation Poisoning[J]. arXiv:2402.08416, 2024.",
        "[8] CHAUDHARI H, SEVERI G, ABASCAL J, et al. Phantom: General Trigger Attacks on Retrieval Augmented Language Generation[J]. arXiv:2405.20485, 2024.",
        "[9] LIU Y, YUAN Z, TIE G, et al. Poisoned-MRAG: Knowledge Poisoning Attacks to Multimodal Retrieval Augmented Generation[J]. arXiv:2503.06254, 2025.",
        "[10] ZHANG B, XIN H, FANG M, et al. Traceback of Poisoning Attacks to Retrieval-Augmented Generation[C]//The Web Conference. 2025.",
        "[11] XIANG C, WU T, ZHONG Z, et al. Certifiably Robust RAG against Retrieval Corruption[J]. arXiv:2405.15556, 2024.",
        "[12] ZHANG B, XIN H, LI J, et al. Benchmarking Poisoning Attacks against Retrieval-Augmented Generation[J]. arXiv:2505.18543, 2025.",
        "[13] SU J, ZHOU J P, ZHANG Z, et al. Towards More Robust Retrieval-Augmented Generation: Evaluating RAG Under Adversarial Poisoning Attacks[J]. arXiv:2412.16708, 2024.",
        "[14] MU Y, HU H, LI F, et al. Towards Secure Retrieval-Augmented Generation: A Comprehensive Review of Threats, Defenses and Benchmarks[J]. arXiv:2603.21654, 2026.",
        "[15] AGGARWAL P, MURAHARI V, RAJPUROHIT T, et al. GEO: Generative Engine Optimization[C]//KDD. 2024.",
        "[16] LIU N F, ZHANG T, LIANG P. Evaluating Verifiability in Generative Search Engines[C]//Findings of EMNLP. 2023.",
        "[17] GAO T, YEN H, YU J, CHEN D. Enabling Large Language Models to Generate Text with Citations[C]//EMNLP. 2023.",
        "[18] ES S, JAMES J, ESPINOSA-ANKE L, SCHOCKAERT S. RAGAS: Automated Evaluation of Retrieval Augmented Generation[C]//EACL System Demonstrations. 2024.",
        "[19] SAAD-FALCON J, KHATTAB O, POTTS C, ZAHARIA M. ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems[C]//NAACL. 2024.",
        "[20] MIN S, KRISHNA K, LYU X, et al. FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation[C]//EMNLP. 2023.",
        "[21] THORNE J, VLACHOS A, CHRISTODOULOPOULOS C, MITTAL A. FEVER: A Large-scale Dataset for Fact Extraction and Verification[C]//NAACL. 2018.",
        "[22] EDGE D, TRINH H, CHENG N, et al. From Local to Global: A Graph RAG Approach to Query-Focused Summarization[J]. arXiv:2404.16130, 2024.",
        "[23] HAN H, WANG Y, SHOMER H, et al. Retrieval-Augmented Generation with Graphs (GraphRAG)[J]. arXiv:2501.00309, 2025.",
        "[24] ZHOU Y, LIU Y, LI X, et al. Trustworthiness in Retrieval-Augmented Generation Systems: A Survey[J]. arXiv:2409.10102, 2024.",
        "[25] NI B, LIU Z, WANG L, et al. Towards Trustworthy Retrieval Augmented Generation for Large Language Models: A Survey[J]. arXiv:2502.06872, 2025.",
        "[26] NIU C, WU Y, ZHU J, et al. RAGTruth: A Hallucination Corpus for Developing Trustworthy Retrieval-Augmented Language Models[J]. arXiv:2401.00396, 2024.",
        "[27] VELIČKOVIĆ P, CUCURULL G, CASANOVA A, et al. Graph Attention Networks[C]//ICLR. 2018.",
        "[28] SCHLICHTKRULL M, KIPF T N, BLOEM P, et al. Modeling Relational Data with Graph Convolutional Networks[C]//ESWC. 2018.",
        "[29] WANG X, JI H, SHI C, et al. Heterogeneous Graph Attention Network[C]//The Web Conference. 2019.",
        "[30] NIST. Zero Trust Architecture: NIST Special Publication 800-207[R]. Gaithersburg: National Institute of Standards and Technology, 2020.",
        "[31] CEMRI M C, PAN M Z, YANG S, et al. Why Do Multi-Agent LLM Systems Fail?[J]. arXiv:2503.13657, 2025.",
        "[32] HE P, LIN Y, DONG S, et al. Red-Teaming LLM Multi-Agent Systems via Communication Attacks[J]. arXiv:2502.14847, 2025.",
        "[33] LEE D, TIWARI M. Prompt Infection: LLM-to-LLM Prompt Injection within Multi-Agent Systems[J]. arXiv:2410.07283, 2024.",
        "[34] WANG S, ZHANG G, YU M, et al. G-Safeguard: A Topology-Guided Security Lens and Treatment on LLM-based Multi-agent Systems[C]//ACL. 2025: 7261-7276.",
        "[35] JAMSHIDI S, MORADI DAKHEL A, NAFI K W, KHOMH F. Hallucination Cascade: Analyzing Error Propagation in Multi-Agent LLM Systems[J]. arXiv:2606.07937, 2026.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0)
        p.paragraph_format.hanging_indent = Pt(24)
        set_run_font(p.add_run(ref), size=10.5)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def build_two_block_report(figures: Sequence[Path]) -> None:
    """按参考作品报告的五章结构生成统一融合报告。"""
    (
        fig_background,
        fig_overall,
        fig_framework,
        fig_provenance,
        fig_correction,
        fig_consensusguard,
        fig_zhiyuan,
        fig_dualrisk_threshold,
        fig_dualrisk_lambda,
        fig_trust,
        fig_agent_robustness,
        fig_agent_threshold,
    ) = figures
    experiment = json.loads(EXPERIMENT_RESULT.read_text(encoding="utf-8"))
    rag_result = experiment["rag"]
    agent_rows = experiment["multi_agent"]["rows"]
    doc = Document()
    setup_document(doc)
    add_cover(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("填写说明"), size=16, bold=True)
    add_list(
        doc,
        [
            "本报告围绕作品设计、实现原理、测试方案、创新点和应用价值展开，内容与演示系统保持一致。",
            "作品融合 ConsensusGuard 的多智能体级联错误治理能力与智源净域的 RAG 知识投毒因果验证能力，"
            "统一形成证据—声明—消息—动作可信链。",
            "表格统一采用三线表，正文、标题、图注和表注统一使用宋体。",
            "公开数据集仅用于防御评测；本地投毒样本仅用于封闭环境中的检测演示，不连接真实服务。",
            "报告中的性能数据须由实际程序运行产生。未完成实测的项目以“待实测”标记，不使用虚构结果。",
        ],
        numbered=True,
    )
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("目  录"), size=16, bold=True)
    add_toc(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    restart_page_number(body_section, 1)

    doc.add_heading("摘  要", level=1)
    add_body(
        doc,
        "随着大语言模型从单体问答扩展为多智能体协同系统，规划、检索、验证和执行被分配给不同 Agent。"
        "单个 Agent 的幻觉、错误工具返回、身份冒用或通信篡改，可能沿消息依赖链持续放大并形成错误共识。"
        "系统吸收 ConsensusGuard 的零信任声明包、声明依赖图、级联错误检测、证据加权共识和状态回滚机制，"
        "验证每一条关键声明的身份、权限、证据与传播关系，实现从异常发现到可信纠偏的闭环。"
    )
    add_body(
        doc,
        "检索增强生成（RAG）依赖外部文档和知识库，攻击者可能通过定向植入、站群转载或伪造权威来源，使污染"
        "片段被高频检索并诱导生成错误答案。系统同时吸收智源净域的统一 Evidence 对象，联合检索吸附性 RAS、答案"
        "诱导性 GIS 和 DualRisk 进行候选筛查，再通过四路反事实和 CausalScore 验证真正致错的知识片段，"
        "并利用异构传播图、声明—证据矩阵、风险隔离、风险感知重检索和可信重生成完成溯源与恢复。"
    )
    add_body(
        doc,
        "作品以“证据—声明—消息—动作”为统一分析主线，将多 Agent 内部的默认信任、级联错误与 RAG 外部知识"
        "污染放入同一因果链处理。系统采用 FastAPI、TF-IDF、启发式 NLI 和 NetworkX 构建可复现实验原型，"
        "所有风险案例均在本地模拟环境运行，适用于信息安全作品赛展示、企业知识库审计和智能体工作流安全验证。"
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(
        p.add_run(
            "关键词：多智能体系统；零信任协同；级联错误；检索增强生成；知识投毒；因果验证；可信纠偏"
        ),
        bold=True,
    )
    doc.add_page_break()

    doc.add_heading("第一章 作品概述", level=1)
    doc.add_heading("1.1 背景分析", level=2)
    add_body(
        doc,
        "人工智能安全已从单一模型可靠性问题上升为数字中国建设、国家网络空间安全和新质生产力安全发展的基础"
        "保障。《新一代人工智能发展规划》提出建立安全可控的人工智能技术体系；《全球人工智能治理倡议》强调"
        "发展与安全并重、提升人工智能技术的安全性、可靠性、可控性和公平性。随着生成式人工智能和智能体进入"
        "政务、工业、金融与安全运营，外部知识污染和内部协同错误可能放大为错误决策链，直接影响关键业务连续性。"
    )
    add_body(
        doc,
        "从法律与监管角度，《网络安全法》《数据安全法》《个人信息保护法》和《网络数据安全管理条例》对网络"
        "运行安全、数据处理安全和个人信息权益保护提出明确要求；《生成式人工智能服务管理暂行办法》要求提高"
        "生成式人工智能服务的透明度以及生成内容的准确性和可靠性；《人工智能生成合成内容标识办法》进一步"
        "强化生成内容全流程治理。RAG 知识库和多 Agent 决策链若缺少证据审计、权限控制、来源追踪和纠偏机制，"
        "将难以满足安全可控、责任可追溯和风险可处置的治理要求。"
    )
    add_body(
        doc,
        "多智能体系统通过角色分工提高复杂任务的完成能力，但其通信通常以自然语言消息为主。若接收方直接采纳"
        "上游结论而不验证身份、权限和证据，局部幻觉会经任务规划、交叉验证和执行环节逐级扩散。多数投票也不能"
        "天然解决问题，因为多个 Agent 可能共享同一污染证据，或被同一个异常节点影响。"
    )
    add_body(
        doc,
        "RAG 将外部知识引入模型上下文，缓解参数知识陈旧问题，但同时扩大了供应链攻击面。污染文档可以利用"
        "关键词覆盖、语义贴合和来源伪装获得异常检索优势；一旦进入 Top-K，其错误结论可能被答案引用并继续传递"
        "给下游 Agent。仅做静态关键词过滤无法判断某一片段是否真正导致错误。"
    )
    add_figure(doc, fig_background, 1, "作品面向的两类信息污染入口与治理闭环")

    doc.add_heading("1.2 相关工作", level=2)
    doc.add_heading("1.2.1 国家政策、法律规范与人工智能安全治理研究", level=3)
    add_body(
        doc,
        "NIST SP 800-207 提出的持续验证、最小权限和动态访问决策，为多智能体零信任协同提供治理基础。"
        "在多 Agent 自动执行逐步进入审批、运维和安全分析场景后，身份可验证、职责可约束、结论可审计和故障"
        "可恢复成为系统设计的基本要求。"
    )
    doc.add_heading("1.2.2 多 Agent 失败机理、通信攻击与级联错误研究", level=3)
    add_body(
        doc,
        "MAST-Data 对多智能体失败进行细粒度归因，Hallucination Cascade 关注错误传播，Prompt Infection、"
        "多智能体通信攻击和 G-Safeguard 分别从消息注入、拓扑攻击与图安全角度揭示协作风险。这些研究说明，"
        "检测对象应从最终答案前移至 Agent 声明、通信边和任务依赖。"
    )
    doc.add_heading("1.2.3 RAG 基础、知识投毒攻击与必要条件研究", level=3)
    add_body(
        doc,
        "RAG、DPR、Natural Questions、HotpotQA 和 MS MARCO 奠定了检索与生成评测基础。PoisonedRAG、"
        "Pandora、Phantom 与 Poisoned-MRAG 证明攻击者可以通过少量定向文档影响检索和生成结果，说明防御必须"
        "同时观察检索异常、答案依赖和传播路径。"
    )
    doc.add_heading("1.2.4 RAG 投毒防御、因果溯源与可信重生成研究", level=3)
    add_body(
        doc,
        "Traceback、RobustRAG 等工作研究投毒溯源与鲁棒检索；ALCE、RAGAS、ARES、FActScore、FEVER 和"
        "RAGTruth 提供引用、忠实度和事实性评测方法。本作品在这些研究基础上，将候选风险检测、反事实因果确认、"
        "图谱溯源和可信重生成组合为可运行流程。"
    )
    doc.add_heading("1.2.5 拜占庭鲁棒共识、事务回滚与可信评价研究", level=3)
    add_body(
        doc,
        "Rethinking the Reliability of Multi-agent System 与 Robust Multi-Agent LLMs under Byzantine Faults"
        "从拜占庭容错角度研究故障节点条件下的鲁棒聚合；SagaLLM 将上下文验证、事务语义和回滚引入多 Agent"
        "规划。RAGAS、ARES、FActScore 与 RAGTruth 则从忠实度、事实性和幻觉标注角度建立可信评价基础。"
        "本作品将上述共识、回滚和评价思想统一到 Claim—Evidence 因果链。"
    )

    doc.add_heading("1.3 作品特色", level=2)
    doc.add_heading("1.3.1 多源安全能力统一证据链的构建", level=3)
    add_table(
        doc,
        1,
        "统一可信链的关键环节与核心能力",
        ["环节", "主要对象", "关键问题", "核心能力", "输出"],
        [
            ["主体与声明接入", "Agent、Claim、消息、动作", "默认信任与级联错误", "零信任校验、依赖记录、权限约束", "可信声明对象"],
            ["知识与证据接入", "Document、Chunk、Evidence、Answer", "RAG 知识投毒", "双条件检测、反事实因果", "可信证据对象"],
            ["联合溯源与纠偏", "证据—声明—消息—动作", "跨层污染传播", "统一标识、风险汇聚、联动隔离与回滚", "风险报告与可信结果"],
        ],
    )
    add_body(
        doc,
        "作品通过 Evidence ID、Claim ID、父子依赖、Agent ID 和 Checkpoint ID 建立"
        "统一关联。RAG 板块发现的高风险证据可触发 ConsensusGuard 重新计算声明可信度；Agent 板块识别的异常"
        "声明也可反向定位其引用的污染 Chunk。"
    )
    doc.add_heading("1.3.2 零信任协同与知识投毒因果验证融合", level=3)
    add_body(
        doc,
        "零信任声明包约束 Agent 身份、权限和消息完整性，RAS/GIS 与四路反事实验证引用证据是否具有投毒因果"
        "贡献，两类结果共同进入传播图和 TrustScore。由此避免只治理内部 Agent 或只过滤外部知识的割裂方案。"
    )
    doc.add_heading("1.3.3 检测与因果验证深度融合，增强纠偏鲁棒性", level=3)
    add_body(
        doc,
        "作品将统计异常检测与反事实因果验证结合：先用级联风险、RAS、GIS 和 DualRisk 缩小候选范围，再通过"
        "声明依赖复核和四路反事实确认根因，最后执行隔离、回滚、重检索与重生成，降低单一阈值误判的影响。"
    )
    doc.add_heading("1.4 应用前景分析", level=2)
    doc.add_heading("1.4.1 脱离在线模型依赖，软计算检测更具部署可行性", level=3)
    add_body(
        doc,
        "系统主要依靠消息元数据、证据关系、文本相似度和图结构进行检测，可在封闭网络和普通 CPU 环境部署。"
        "这使其适用于对数据出域敏感的企业知识库、内部审批和安全运营场景。"
    )
    doc.add_heading("1.4.2 面向多行业智能体与 RAG 治理，具备平台集成价值", level=3)
    add_body(
        doc,
        "系统可用于企业多智能体审批、运维协同、安全分析和自动化决策流程，也可部署在企业 RAG 知识助手、"
        "制度查询、漏洞知识库与合规问答场景。模块化接口便于与现有身份系统、检索服务和审计平台集成。"
    )

    doc.add_heading("第二章 作品设计与实现", level=1)
    doc.add_heading("2.1 系统方案", level=2)
    add_body(
        doc,
        "系统不是两个独立子系统的机械拼接，而是围绕同一条可信链运行：多 Agent 在规划、检索和验证过程中产生"
        "原子声明，声明引用 RAG 返回的 Evidence 与 Chunk；零信任校验负责约束声明主体和权限，知识投毒检测"
        "负责判断引用证据是否异常，联合传播图进一步确认污染如何从知识源进入声明并扩散到答案或动作。"
    )
    doc.add_heading("2.1.1 多 Agent 零信任协同与声明可信链模块设计", level=3)
    doc.add_heading("2.1.1.1 基础知识：零信任、原子声明与拜占庭故障", level=4)
    add_body(
        doc,
        "关键消息被拆分为原子 Claim，并绑定 agent_id、role、permission、evidence_ids、parent_claim_ids、"
        "timestamp、nonce 和 signature。接收方在使用声明前完成身份、角色权限、证据存在性和重放检查。"
    )
    doc.add_heading("2.1.1.2 方法设计：声明依赖图与级联检测", level=4)
    add_body(
        doc,
        "系统记录 Claim 之间的引用、支持、矛盾和派生关系。当上游声明被判定为高风险时，沿依赖边计算受影响"
        "范围，并结合证据一致性、来源独立性和异常传播深度识别级联错误，而不是只检查最终答案。"
    )
    doc.add_heading("2.1.1.3 本文模型：证据加权零信任可信共识", level=4)
    add_body(
        doc,
        "共识权重由证据支持率、来源独立性、角色权限、历史可靠度和当前风险共同决定。共享同一证据源的多个 Agent"
        "不会被视为多个独立支持者，从而降低污染证据通过重复转述形成伪多数的风险。"
    )
    doc.add_heading("2.1.1.4 模型输出：隔离、回滚与重规划", level=4)
    add_body(
        doc,
        "检测到异常后，系统隔离高风险 Agent 或 Claim，回滚到最近可信 Checkpoint，删除受污染的下游状态，"
        "再由剩余可信 Agent 基于独立证据重新规划。全过程生成可审计事件记录。"
    )
    add_figure(doc, fig_consensusguard, 2, "多 Agent 零信任声明、级联检测与可信纠偏流程")

    doc.add_heading("2.1.2 RAG 知识投毒因果验证与可信纠偏模块设计", level=3)
    doc.add_heading("2.1.2.1 基础知识：统一证据对象与 Chunk 管理", level=4)
    add_body(
        doc,
        "系统将网页、企业文档和工具返回统一为 Evidence，对内容进行分块、去重和来源标记，并保留 document_id、"
        "chunk_id、source_id、标签、检索历史和风险评分，保证后续检测与溯源使用同一数据契约。"
    )
    doc.add_heading("2.1.2.2 方法设计：RAS、GIS 与 DualRisk 候选筛查", level=4)
    add_body(
        doc,
        "RAS 衡量 Chunk 是否被异常频繁检索，GIS 衡量答案对该 Chunk 的依赖程度，DualRisk 将检索层异常与"
        "生成层诱导联合起来。只有两个条件同时较高的片段才进入重点因果验证，减少对正常热门知识的误报。"
    )
    doc.add_heading("2.1.2.3 方法设计：四路反事实与因果确认", level=4)
    add_body(
        doc,
        "对候选片段分别执行原始 Top-K、删除可疑片段、仅保留可疑片段、以可信证据替代四路生成。若原答案与"
        "仅可疑答案相近，而删除或替代后结论显著变化，则可疑片段具有更强的致错因果贡献。"
    )
    doc.add_heading("2.1.2.4 本文模型：异构传播图与声明—证据矩阵", level=4)
    add_body(
        doc,
        "图谱覆盖 Page、Document、Chunk、Query、Claim 和 Answer 节点，并记录 contains、retrieved_by、"
        "supports、contradicts、copied_from、caused_error 等关系。声明—证据矩阵使用支持、矛盾和中立三分类"
        "呈现证据结构，辅助定位站群转载和单一污染源。"
    )
    doc.add_heading("2.1.2.5 模型输出：隔离、风险感知重检索与可信重生成", level=4)
    add_body(
        doc,
        "因果确认后的高风险 Chunk 被加入隔离集合，重检索阶段排除其内容及高相似副本，并优先选择独立可信来源。"
        "重生成答案只使用通过验证的证据，对证据不足的声明明确标记不确定性。"
    )
    add_figure(doc, fig_zhiyuan, 3, "RAG 证据接入、知识投毒因果验证与可信重生成流程")

    add_body(
        doc,
        "上述模块共享证据与声明标识：Agent 的每条关键结论都能追溯到 RAG Chunk；Chunk 风险变化会触发关联"
        "Claim 和下游动作重评；异常声明也能反向定位其引用知识。两类检测结果在同一风险对象集合中汇聚，形成"
        "双向联动而非前后割裂的串联。"
    )
    add_figure(doc, fig_overall, 4, "零信任协同与 RAG 投毒因果验证融合总体流程")

    doc.add_heading("2.2 实现原理", level=2)
    doc.add_heading("2.2.1 声明级信息污染识别与证据处理", level=3)
    doc.add_heading("2.2.1.1 声明与污染证据识别", level=4)
    add_table(
        doc,
        2,
        "ConsensusGuard 核心检查项",
        ["检查阶段", "输入", "判断依据", "处置"],
        [
            ["身份与权限", "Agent、角色、操作", "身份有效性、最小权限", "拒绝或降权"],
            ["声明与证据", "Claim、Evidence", "支持关系、来源独立性", "标记分歧或补证"],
            ["传播依赖", "Claim DAG", "父子风险、传播深度", "定位受影响节点"],
            ["可信共识", "多 Agent 结论", "证据权重与历史可靠度", "选取可信结论"],
            ["恢复", "Checkpoint、任务状态", "最近可信状态", "隔离、回滚、重规划"],
        ],
    )
    add_body(
        doc,
        "系统采用“先验证后使用”的消息消费规则。任何缺少证据、权限不足或依赖高风险父声明的结论都不能直接"
        "进入执行阶段；任何被声明引用的 Chunk 也必须同时接受来源、检索吸附性和答案诱导性检查。正常分歧与"
        "恶意协同通过证据独立性和传播拓扑区分，避免把少数但有强证据的正确 Agent 误判。"
    )
    doc.add_heading("2.2.1.2 可信证据处理与风险隔离", level=4)
    add_body(
        doc,
        "通过校验的 Evidence 被写入声明—证据矩阵；存在身份异常、来源复制、内容矛盾或高因果风险的对象进入"
        "隔离集合。隔离以 Chunk、Claim 和 Agent 为不同粒度执行，既阻止污染继续传播，也保留原始对象用于"
        "审计、反事实复核和根因解释。"
    )

    doc.add_heading("2.2.2 级联错误与知识投毒联合检测", level=3)
    doc.add_heading("2.2.2.1 多 Agent 声明依赖与 RAG 证据关联基础", level=4)
    add_body(
        doc,
        "系统把 Agent 消息拆为 Claim，并以 evidence_ids 连接检索 Chunk，以 parent_claim_ids 连接上游声明。"
        "因此，一个最终错误可以同时沿消息依赖边和知识引用边追踪，判断其根因是 Agent 自身幻觉、通信异常，"
        "还是外部污染证据进入上下文后诱导生成。"
    )
    doc.add_heading("2.2.2.2 传统单点检测与多数共识方法局限", level=4)
    add_body(
        doc,
        "仅检查最终答案会丢失传播过程，仅按关键词过滤无法确认致错因果，仅按 Agent 数量投票又可能把同源污染"
        "的重复转述误认为独立共识。系统因此联合主体权限、证据独立性、检索异常、答案依赖与声明传播深度进行"
        "判断，并将候选风险交给反事实验证。"
    )
    doc.add_heading("2.2.2.3 本文模型", level=4)
    add_formula(doc, "RAS(chunk_i) = (freq_i / total_retrievals) / (1 / total_chunks)")
    add_formula(doc, "GIS(chunk_i, answer) = sim(answer, chunk_i) / max_sim(answer, all_chunks)")
    add_formula(doc, "DualRisk(chunk_i) = RAS(chunk_i) × GIS(chunk_i)")
    add_formula(doc, "CausalScore(chunk_i) = 1 - sim(A_orig, A_remove) / sim(A_orig, A_only_suspect)")
    add_formula(
        doc,
        "TrustScore = 0.3×EvidenceSupportRate + 0.2×SourceIndependenceScore "
        "+ 0.3×(1-NormalizedDualRisk) + 0.2×(1-NormalizedCausalScore)",
    )
    add_body(
        doc,
        "上述风险分数与 Agent 身份权限、Claim 父子风险和证据独立性共同进入联合判定。MVP 使用 TF-IDF 余弦"
        "相似度实现 GIS 与反事实答案比较，使用启发式 NLI 判断支持、矛盾和中立关系。各服务通过抽象基类保留"
        "替换接口，后续可接入 BGE、专业 NLI、FAISS/Milvus 和约束大模型生成。"
    )

    doc.add_heading("2.3 基于声明—证据关联图的信息污染溯源实现", level=2)
    doc.add_heading("2.3.1 多 Agent 级联错误传播建模", level=3)
    add_body(
        doc,
        "每条 Claim 记录生成 Agent、父声明、引用证据、接收者和后续动作。系统以有向无环依赖图表示一次任务中"
        "的推理和协作过程，并在发现异常后沿后继边计算受影响声明与动作，沿前驱边搜索最早异常节点。"
    )
    doc.add_heading("2.3.2 证据与风险分数归一化处理", level=3)
    add_body(
        doc,
        "不同检测量纲先归一化到 [0,1]。RAS 使用语料规模和检索总次数校正，GIS 以当前候选最大相似度归一化，"
        "CausalScore 对异常边界进行截断；Agent 历史可靠度、权限有效性和来源独立性同样转换为可组合权重。"
    )
    doc.add_heading("2.3.3 基于异构传播图的跨层依赖建模", level=3)
    doc.add_heading("2.3.3.1 输入与输出结构", level=4)
    add_body(
        doc,
        "输入包括 Page、Document、Chunk、Evidence、Query、Agent、Claim、Answer、Action 与 Checkpoint；"
        "输出包括根因节点排序、污染传播路径、受影响对象集合、路径风险和建议隔离点。"
    )
    doc.add_heading("2.3.3.2 传播贡献计算", level=4)
    add_body(
        doc,
        "每条边根据关系类型设置基础权重，再结合上游节点风险、证据支持或矛盾关系、时间顺序和反事实变化计算"
        "传播贡献。对 copied_from 和 same_claim 关系执行来源聚类，防止站群转载被重复计权。"
    )
    doc.add_heading("2.3.3.3 多路径归因机制", level=4)
    add_body(
        doc,
        "当错误答案同时受到多个 Chunk 和多个 Agent 影响时，系统枚举关键有向路径并合并共享前缀，分别报告"
        "知识源贡献、声明传播贡献和执行放大贡献，从而避免把所有风险简单归于最后一个生成 Agent。"
    )
    doc.add_heading("2.3.4 本文模型", level=3)
    add_body(
        doc,
        "联合图在 RAG 图谱基础上扩展 Agent、Message、Claim、Action 与 Checkpoint 节点。溯源从错误 Answer"
        "或高风险 Action 反向遍历 caused_error、derived_from、supports、retrieved_by 和 contains 等边，"
        "得到污染源、传播路径、受影响声明和最终动作。"
    )
    add_figure(doc, fig_provenance, 5, "证据—声明—Agent—动作联合溯源关系")

    doc.add_heading("2.4 基于可信度更新的动态纠偏", level=2)
    doc.add_heading("2.4.1 本文算法", level=3)
    add_body(
        doc,
        "纠偏按风险对象分层执行：知识层隔离 Chunk 和相似副本；声明层撤销被污染证据支持的 Claim；协同层隔离"
        "异常 Agent 或降低其权重；状态层回滚到可信 Checkpoint；生成层以独立可信证据重新检索和生成。复核未通过"
        "时继续补证，不允许未经验证的结果恢复执行。"
    )
    add_figure(doc, fig_correction, 6, "分层隔离、回滚、重检索与可信重生成闭环")

    doc.add_heading("2.5 实时证据比对与污染判断", level=2)
    doc.add_heading("2.5.1 声明证据与实时检索结果对比", level=3)
    add_body(
        doc,
        "系统对关键 Claim 的引用证据与隔离后的重检索结果进行比对。当独立可信来源对原声明形成矛盾，或删除"
        "可疑证据后答案发生实质变化时，提高声明风险并冻结其下游动作；若补充证据保持一致，则降低误报风险。"
    )
    doc.add_heading("2.5.2 检索吸附性与答案诱导集中度", level=3)
    add_body(
        doc,
        "RAS 反映检索结果是否过度集中到特定 Chunk，GIS 反映答案是否过度贴近某一证据。两者与 Agent 间的"
        "转述集中度联合观察，可识别单一污染源被多个节点重复引用而形成的伪共识。"
    )
    doc.add_heading("2.5.3 综合信息污染风险评分函数", level=3)
    add_body(
        doc,
        "综合风险由主体可信度、证据独立性、DualRisk、CausalScore、NLI 矛盾率和传播深度构成。评分只用于"
        "排序与处置触发，最终高风险结论仍需反事实结果或明确的权限、签名异常作为可解释依据。"
    )
    doc.add_heading("2.5.4 算法思路与实现：零信任协同—因果验证—可信纠偏算法", level=3)
    add_body(
        doc,
        "系统后端采用 Python 3.10+ 与 FastAPI，服务层按检测能力拆分；NetworkX 维护异构传播图；"
        "scikit-learn 提供 TF-IDF 相似度；Pydantic v2 统一数据模型。MVP 不依赖 GPU 和在线模型接口，便于"
        "在比赛环境中复现。"
    )
    add_figure(doc, fig_framework, 7, "系统分层实现架构")
    doc.add_page_break()
    add_table(
        doc,
        3,
        "系统主要功能模块",
        ["层级", "主体与声明能力", "知识与证据能力", "融合输出"],
        [
            ["数据对象", "Agent、Claim、Message、Checkpoint", "Evidence、Document、Chunk、Answer", "统一标识与依赖"],
            ["检测", "身份权限、级联错误、异常协同", "RAS、GIS、DualRisk", "风险对象集合"],
            ["验证", "证据加权共识、传播复核", "四路反事实、CausalScore", "因果与可信评分"],
            ["溯源", "Claim DAG、Agent 通信图", "异构投毒传播图", "污染路径"],
            ["处置", "隔离、回滚、重规划", "Chunk 隔离、重检索、重生成", "纠偏结果与报告"],
        ],
    )

    doc.add_heading("2.6 性能指标", level=2)
    add_table(
        doc,
        4,
        "拟评测指标及计算口径",
        ["板块", "指标", "计算口径", "状态"],
        [
            ["协同与声明层", "任务正确率", "不同异常 Agent 比例下正确任务结果比例", "本地仿真实测"],
            ["联合溯源层", "源头定位率", "真实异常 Agent 是否被风险集合命中", "本地仿真实测"],
            ["可信纠偏层", "纠偏成功率", "基线错误而融合方法恢复正确的比例", "本地仿真实测"],
            ["知识与证据层", "投毒检测 F1", "高风险 Chunk 检测精确率与召回率", "本地案例实测"],
            ["可信纠偏层", "TrustScore 提升", "隔离与重生成前后的可信评分差值", "本地案例实测"],
            ["生成与复核层", "投毒证据隔离率", "被隔离投毒证据占投毒证据总数", "本地案例实测"],
            ["联合系统", "端到端时延", "从查询输入到风险报告输出的耗时", "待实测"],
        ],
    )

    doc.add_heading("第三章 作品测试与分析", level=1)
    doc.add_heading("3.1 测试环境", level=2)
    add_table(
        doc,
        5,
        "MVP 测试环境",
        ["类别", "配置"],
        [
            ["操作系统", "Windows 11 / Linux，均可本地运行"],
            ["运行环境", "Python 3.10+，CPU 模式"],
            ["核心依赖", "FastAPI、Pydantic v2、NetworkX、scikit-learn、NumPy、pytest"],
            ["模型与服务", "不调用在线 API，不要求 GPU，不进行大模型训练"],
            ["数据安全", "公开数据集离线评测与本地模拟案例，不访问真实业务服务"],
        ],
    )

    doc.add_heading("3.2 测试数据集", level=2)
    doc.add_heading("3.2.1 公开数据集", level=3)
    doc.add_heading("3.2.1.1 MAST-Data、MultiAgentBench 与 AgentDojo 数据集", level=4)
    add_body(
        doc,
        "MAST-Data 提供多智能体系统失败轨迹和细粒度分类，可用于声明级错误定位与级联传播分析。"
        "MultiAgentBench 提供多智能体协作任务和评测框架，可用于构造正常协作、分歧和异常通信对照；"
        "AgentDojo 面向工具型智能体安全评测，可用于验证角色权限、工具调用和不可信输入隔离。"
    )
    doc.add_heading("3.2.1.2 PoisonedRAG 与 RAGTruth 数据集", level=4)
    add_body(
        doc,
        "PoisonedRAG 提供知识投毒攻击与评测设置，可用于检测和因果验证复现；RAGTruth 提供 RAG 输出中的"
        "幻觉标注，可作为声明—证据一致性与可信重生成质量的辅助评测集。"
    )
    doc.add_heading("3.2.1.3 ALCE 与 CFEVER 数据集", level=4)
    add_body(
        doc,
        "ALCE 适合评估答案引用完整性和可验证性；CFEVER 提供中文事实验证样本，可用于中文支持、矛盾和中立"
        "判断。二者用于评价纠偏后答案，不作为知识投毒攻击数据。"
    )
    add_table(
        doc,
        6,
        "建议采用的公开数据集与链接",
        ["板块", "数据集", "用途", "公开链接"],
        [
            ["ConsensusGuard", "MAST-Data", "多智能体失败归因与级联错误", "https://arxiv.org/abs/2503.13657"],
            ["ConsensusGuard", "MultiAgentBench", "多智能体协作与通信评测", "https://github.com/MultiagentBench/MARBLE"],
            ["ConsensusGuard", "AgentDojo", "工具调用与权限安全评测", "https://github.com/ethz-spylab/agentdojo"],
            ["智源净域", "PoisonedRAG", "RAG 知识投毒检测与因果验证", "https://arxiv.org/abs/2402.07867"],
            ["智源净域", "RAGTruth", "RAG 幻觉和忠实度评测", "https://arxiv.org/abs/2401.00396"],
            ["智源净域", "ALCE", "引用正确性与可验证性", "https://github.com/princeton-nlp/ALCE"],
            ["智源净域", "CFEVER", "中文事实验证", "https://ikmlab.github.io/CFEVER/"],
        ],
    )
    add_body(
        doc,
        "在公开数据集之外，本地受控案例覆盖企业制度投毒、漏洞状态投毒、安全认证投毒、站群投毒和良性错误负例，并扩展身份冒用、"
        "权限越界、单节点幻觉、错误共识和共享污染证据等多智能体情形。每个案例包含真实标签、正确答案、攻击"
        "目标、文档来源、Claim 依赖和期望处置，且带有仅限本地防御演示的安全声明。"
    )

    doc.add_heading("3.3 测试方案", level=2)
    doc.add_heading("3.3.1 多 Agent 声明与级联错误数据集划分", level=3)
    add_list(
        doc,
        [
            "正常协作：各 Agent 使用独立可信证据，验证系统不会干扰正确任务。",
            "单点级联：向上游 Agent 注入本地模拟错误声明，观察根因定位和下游影响范围。",
            "伪多数：多个 Agent 转述同一错误证据，验证来源独立性去重是否抑制错误共识。",
            "权限异常：模拟角色越权和重放消息，验证零信任声明包的拒绝与审计。",
            "恢复测试：隔离异常节点并回滚 Checkpoint，验证任务能否重新规划并完成。",
        ],
        numbered=True,
    )
    doc.add_heading("3.3.2 RAG 知识投毒与可信纠偏数据集划分", level=3)
    add_list(
        doc,
        [
            "在不同投毒比例和 Top-K 下计算 RAS、GIS、DualRisk 与检测 F1。",
            "对候选 Chunk 执行四路反事实，比较因果分与真实投毒标签的一致性。",
            "构建站群转载和独立来源对照，检查传播图是否定位原始污染文档。",
            "隔离高风险 Chunk 后重新检索，比较攻击成功率、正确答案保持率和引用质量。",
            "对良性过时信息、正常热门文档和证据不足样本进行误报测试。",
        ],
        numbered=True,
    )
    doc.add_heading("3.3.3 评估要求", level=3)
    add_body(
        doc,
        "数据按案例或来源划分训练、验证与测试集合，避免同源文档泄漏。固定随机种子和依赖版本，记录每次运行的"
        "配置、原始输出和指标脚本。公开数据遵守原许可证，本地模拟数据不发布到互联网。"
    )
    doc.add_heading("3.3.4 测试结果", level=3)
    doc.add_heading("3.3.4.1 本地受控实验结果", level=4)
    best_rag = rag_result["best_baselines"]
    add_table(
        doc,
        7,
        "RAG 投毒检测方法实测对比",
        ["方法", "最优阈值", "Precision", "Recall", "F1", "FPR"],
        [
            [
                name,
                f'{row["threshold"]:.2f}',
                f'{row["precision"]:.3f}',
                f'{row["recall"]:.3f}',
                f'{row["f1"]:.3f}',
                f'{row["fpr"]:.3f}',
            ]
            for name, row in best_rag.items()
        ],
    )
    add_body(
        doc,
        f"该实验在两个本地受控案例的 {rag_result['sample_count']} 条 Top-K 证据上运行，其中投毒证据 "
        f"{rag_result['poison_count']} 条。DualRisk 在阈值 0.60 时取得 Precision=0.800、Recall=1.000、"
        "F1=0.889。样本规模较小，因此该结果只说明当前原型在受控案例上的可运行性，不代表公开数据集性能。"
    )
    add_figure(doc, fig_dualrisk_threshold, 8, "DualRisk 阈值对 Precision、Recall 与 F1 的影响")
    add_figure(doc, fig_dualrisk_lambda, 9, "DualRisk 几何项权重 λ 参数实验")

    add_table(
        doc,
        8,
        "可信纠偏前后 TrustScore 实测结果",
        ["案例", "纠偏前", "纠偏后", "提升", "隔离投毒证据"],
        [
            [
                "企业制度投毒" if row["case_id"] == "case_enterprise_rag_poisoning" else "站群投毒",
                f'{row["trust_before"]:.2f}',
                f'{row["trust_after"]:.2f}',
                f'{row["trust_gain"]:.2f}',
                f'{row["isolated_count"]}/{row["poison_count"]}',
            ]
            for row in rag_result["case_results"]
        ],
    )
    add_figure(doc, fig_trust, 10, "风险隔离与可信重生成前后的 TrustScore")
    doc.add_page_break()

    selected_agents = [
        row for row in agent_rows
        if row["risk_threshold"] == 0.55 and row["malicious_ratio"] in (0.14, 0.29, 0.43)
    ]
    add_table(
        doc,
        9,
        "多 Agent 级联错误受控仿真实测结果",
        ["异常 Agent 比例", "无防护", "多数投票", "单验证 Agent", "本文方法", "纠偏成功率"],
        [
            [
                f'{row["malicious_ratio"]:.2f}',
                f'{row["no_defense_accuracy"]:.3f}',
                f'{row["majority_accuracy"]:.3f}',
                f'{row["single_verifier_accuracy"]:.3f}',
                f'{row["proposed_accuracy"]:.3f}',
                f'{row["correction_rate"]:.3f}',
            ]
            for row in selected_agents
        ],
    )
    add_body(
        doc,
        "多 Agent 实验使用 7 个 Agent、固定随机种子 20260614，每个异常比例与风险阈值组合重复 500 次。"
        "本文方法将签名有效性、证据支持度、来源独立性和 Claim 风险共同用于加权共识。当异常 Agent 比例为 "
        "0.43 时，本文方法任务正确率为 0.924，高于无防护的 0.456、多数投票的 0.408 和单验证 Agent 的 "
        "0.764。该实验为固定规则受控仿真，后续仍需在 MAST-Data 与 MultiAgentBench 上验证外部有效性。"
    )
    add_figure(doc, fig_agent_robustness, 11, "不同异常 Agent 比例下各方法任务正确率")
    add_figure(doc, fig_agent_threshold, 12, "联合风险阈值对任务正确率与源头定位率的影响")

    doc.add_heading("3.3.5 模型对比", level=3)
    add_list(
        doc,
        [
            "ConsensusGuard 对比：多数投票、仅最终答案核验、无来源独立性加权、无回滚机制。",
            "智源净域对比：仅关键词过滤、仅 RAS、仅 GIS、无反事实验证、无相似副本隔离。",
            "联合系统消融：断开 Evidence—Claim 关联，观察跨板块根因定位与纠偏能力变化。",
        ],
        numbered=True,
    )
    doc.add_heading("3.3.6 测试结果分析", level=3)
    add_body(
        doc,
        "结果分析应同时报告安全性与可用性：检测率提升不能以大量误报和任务失败为代价。重点检查共享证据形成"
        "伪多数、正常分歧被误判、热门文档 RAS 偏高、反事实答案为空以及回滚状态不完整等边界情况。"
    )

    doc.add_heading("第四章 创新性说明", level=1)
    doc.add_heading("4.1 面向多 Agent 与 RAG 融合场景的信息污染可信治理模型", level=2)
    add_list(
        doc,
        [
            "将零信任从网络访问扩展到 Agent 声明消费，关键 Claim 必须绑定身份、权限、证据和依赖。",
            "从最终答案检测前移到声明依赖图，定位错误如何跨 Agent 传播并影响动作。",
            "以证据独立性和角色权限构建可信共识，抑制共享污染证据造成的伪多数。",
            "将检测与隔离、回滚、重规划结合，形成可恢复的协同安全闭环。",
        ],
        numbered=True,
    )
    doc.add_heading("4.2 零信任协同与知识投毒因果验证的联合实现", level=2)
    add_list(
        doc,
        [
            "以 RAS 与 GIS 双条件联合识别同时具有检索吸附性和答案诱导性的可疑 Chunk。",
            "通过四路反事实与 CausalScore 区分统计相关和真实致错因果贡献。",
            "联合异构传播图与声明—证据矩阵，定位原始污染源、转载链和受影响答案。",
            "从高风险 Chunk 隔离延伸到相似副本清除、风险感知重检索和可信重生成。",
        ],
        numbered=True,
    )
    doc.add_heading("4.3 面向信息污染传播链的融合式溯源与可信纠偏模型", level=2)
    add_body(
        doc,
        "作品建立 Evidence—Claim—Message—Action 的统一链路，使外部知识污染和内部级联错误能够在同一张"
        "因果传播图中被解释。RAG 风险触发 Agent 声明重评，Agent 异常反向定位引用知识，实现跨板块联动隔离与"
        "可信恢复。"
    )

    doc.add_heading("第五章 总结", level=1)
    doc.add_heading("5.1 工作总结", level=2)
    add_body(
        doc,
        "本作品首先以 ConsensusGuard 治理多智能体协作中的默认信任、错误传播和恢复问题，再以智源净域治理"
        "RAG 外部知识投毒、因果溯源和可信重生成问题。两个板块通过统一证据与声明链形成端到端信息污染防御闭环，"
        "既能说明错误从何而来、如何传播，也能给出隔离、回滚、重检索和重生成后的可信结果。"
    )
    doc.add_heading("5.2 未来展望", level=2)
    add_body(
        doc,
        "后续将在真实可授权数据上完成可复现实测，完善中文多智能体失败标注和 RAG 投毒基准；在保持服务接口"
        "不变的前提下，引入更强的语义向量、NLI、异构图学习和约束生成模型，并加强大规模图谱性能、权限系统对接"
        "与人工复核工作流。"
    )

    doc.add_heading("参考文献", level=1)
    references = [
        "[1] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS. 2020.",
        "[2] KARPUKHIN V, OGUZ B, MIN S, et al. Dense Passage Retrieval for Open-Domain Question Answering[C]//EMNLP. 2020.",
        "[3] KWIATKOWSKI T, PALOMAKI J, REDFIELD O, et al. Natural Questions: A Benchmark for Question Answering Research[J]. TACL, 2019, 7: 452-466.",
        "[4] YANG Z, QI P, ZHANG S, et al. HotpotQA: A Dataset for Diverse, Explainable Multi-hop Question Answering[C]//EMNLP. 2018.",
        "[5] NGUYEN T, ROSENBERG M, SONG X, et al. MS MARCO: A Human Generated Machine Reading Comprehension Dataset[C]//CoCo@NIPS. 2016.",
        "[6] ZOU W, GENG R, WANG B, JIA J. PoisonedRAG: Knowledge Corruption Attacks to Retrieval-Augmented Generation of Large Language Models[C]//USENIX Security Symposium. 2025.",
        "[7] DENG G, LIU Y, WANG K, et al. Pandora: Jailbreak GPTs by Retrieval Augmented Generation Poisoning[J]. arXiv:2402.08416, 2024.",
        "[8] CHAUDHARI H, SEVERI G, ABASCAL J, et al. Phantom: General Trigger Attacks on Retrieval Augmented Language Generation[J]. arXiv:2405.20485, 2024.",
        "[9] LIU Y, YUAN Z, TIE G, et al. Poisoned-MRAG: Knowledge Poisoning Attacks to Multimodal Retrieval Augmented Generation[J]. arXiv:2503.06254, 2025.",
        "[10] ZHANG B, XIN H, FANG M, et al. Traceback of Poisoning Attacks to Retrieval-Augmented Generation[C]//The Web Conference. 2025.",
        "[11] XIANG C, WU T, ZHONG Z, et al. Certifiably Robust RAG against Retrieval Corruption[J]. arXiv:2405.15556, 2024.",
        "[12] ZHANG B, XIN H, LI J, et al. Benchmarking Poisoning Attacks against Retrieval-Augmented Generation[J]. arXiv:2505.18543, 2025.",
        "[13] SU J, ZHOU J P, ZHANG Z, et al. Towards More Robust Retrieval-Augmented Generation: Evaluating RAG Under Adversarial Poisoning Attacks[J]. arXiv:2412.16708, 2024.",
        "[14] GAO T, YEN H, YU J, CHEN D. Enabling Large Language Models to Generate Text with Citations[C]//EMNLP. 2023.",
        "[15] ES S, JAMES J, ESPINOSA-ANKE L, SCHOCKAERT S. RAGAS: Automated Evaluation of Retrieval Augmented Generation[C]//EACL System Demonstrations. 2024.",
        "[16] SAAD-FALCON J, KHATTAB O, POTTS C, ZAHARIA M. ARES: An Automated Evaluation Framework for Retrieval-Augmented Generation Systems[C]//NAACL. 2024.",
        "[17] MIN S, KRISHNA K, LYU X, et al. FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation[C]//EMNLP. 2023.",
        "[18] THORNE J, VLACHOS A, CHRISTODOULOPOULOS C, MITTAL A. FEVER: A Large-scale Dataset for Fact Extraction and Verification[C]//NAACL. 2018.",
        "[19] NIU C, WU Y, ZHU J, et al. RAGTruth: A Hallucination Corpus for Developing Trustworthy Retrieval-Augmented Language Models[J]. arXiv:2401.00396, 2024.",
        "[20] NIST. Zero Trust Architecture: NIST Special Publication 800-207[R]. Gaithersburg: National Institute of Standards and Technology, 2020.",
        "[21] CEMRI M C, PAN M Z, YANG S, et al. Why Do Multi-Agent LLM Systems Fail?[J]. arXiv:2503.13657, 2025.",
        "[22] HE P, LIN Y, DONG S, et al. Red-Teaming LLM Multi-Agent Systems via Communication Attacks[J]. arXiv:2502.14847, 2025.",
        "[23] LEE D, TIWARI M. Prompt Infection: LLM-to-LLM Prompt Injection within Multi-Agent Systems[J]. arXiv:2410.07283, 2024.",
        "[24] WANG S, ZHANG G, YU M, et al. G-Safeguard: A Topology-Guided Security Lens and Treatment on LLM-based Multi-agent Systems[C]//ACL. 2025: 7261-7276.",
        "[25] JAMSHIDI S, MORADI DAKHEL A, NAFI K W, KHOMH F. Hallucination Cascade: Analyzing Error Propagation in Multi-Agent LLM Systems[J]. arXiv:2606.07937, 2026.",
        "[26] 国务院. 新一代人工智能发展规划: 国发〔2017〕35号[Z]. 2017.",
        "[27] 国家互联网信息办公室, 国家发展和改革委员会, 教育部, 等. 生成式人工智能服务管理暂行办法[Z]. 2023.",
        "[28] 国家互联网信息办公室, 工业和信息化部, 公安部, 国家广播电视总局. 人工智能生成合成内容标识办法[Z]. 2025.",
        "[29] 中华人民共和国网络安全法[Z]. 2016, 2025年修正.",
        "[30] 中华人民共和国数据安全法[Z]. 2021.",
        "[31] 中华人民共和国个人信息保护法[Z]. 2021.",
        "[32] 国务院. 网络数据安全管理条例: 国务院令第790号[Z]. 2024.",
        "[33] XIE Y, ZHU C, ZHANG X, et al. From Spark to Fire: Modeling and Mitigating Error Cascades in LLM-Based Multi-Agent Collaboration[J]. arXiv:2603.04474, 2026.",
        "[34] ZHOU J, WANG L, et al. GUARDIAN: Safeguarding LLM Multi-Agent Collaborations with Temporal Graph Modeling[J]. arXiv:2505.19234, 2025.",
        "[35] ZHENG L, CHEN J, YIN Q, et al. Rethinking the Reliability of Multi-agent System: A Perspective from Byzantine Fault Tolerance[J]. arXiv:2511.10400, 2025.",
        "[36] LEE H, YUN V D, OH H, et al. Robust Multi-Agent LLMs under Byzantine Faults[J]. arXiv:2605.09076, 2026.",
        "[37] CHANG E Y, GENG L. SagaLLM: Context Management, Validation, and Transaction Guarantees for Multi-Agent LLM Planning[J]. arXiv:2503.11951, 2025.",
    ]
    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.hanging_indent = Pt(24)
        set_run_font(p.add_run(ref), size=10.5)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    configure_matplotlib()
    figures = [
        figure_background(),
        figure_overall_flow(),
        figure_framework(),
        figure_provenance(),
        figure_correction(),
        figure_consensusguard(),
        figure_zhiyuan_rag(),
        EXPERIMENT_FIGURE_DIR / "exp_dualrisk_threshold.png",
        EXPERIMENT_FIGURE_DIR / "exp_dualrisk_lambda.png",
        EXPERIMENT_FIGURE_DIR / "exp_trust_before_after.png",
        EXPERIMENT_FIGURE_DIR / "exp_multi_agent_robustness.png",
        EXPERIMENT_FIGURE_DIR / "exp_multi_agent_threshold.png",
    ]
    build_two_block_report(figures)
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
