#!/usr/bin/env python3
"""Build the 2.1.2 cluster-aware module design DOCX from markdown."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".docx_deps"))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = ROOT / "docs" / "module_design_2_1_2_cluster.md"
OUT = ROOT / "docs" / "module_design_2_1_2_cluster_project_figures.docx"


def set_run_font(run, size: int = 12) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.size = Pt(size)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.45
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, 12)


def add_image(doc: Document, base: Path, markdown_line: str) -> None:
    caption = markdown_line[2 : markdown_line.index("](")]
    rel_path = markdown_line[markdown_line.index("](") + 2 : -1]
    image_path = base / rel_path
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
    caption_paragraph = doc.add_paragraph(caption)
    caption_paragraph.alignment = 1
    caption_paragraph.paragraph_format.line_spacing = 1.15
    caption_paragraph.paragraph_format.space_after = Pt(6)
    for run in caption_paragraph.runs:
        set_run_font(run, 10)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    base = SRC.parent

    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    for raw_line in SRC.read_text(encoding="utf-8").splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            add_image(doc, base, line)
        else:
            add_body_paragraph(doc, line)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
