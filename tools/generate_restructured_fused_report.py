#!/usr/bin/env python3
"""按指定目录生成“智源净域”全国大学生信息安全作品赛作品报告。"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from typing import Iterable

ROOT = Path(__file__).resolve().parents[2]
sys.path.insert(0, str(ROOT / "zyjd_system" / "tools"))

from generate_fused_report import (  # noqa: E402
    ASSET_DIR,
    COLORS,
    EXPERIMENT_FIGURE_DIR,
    EXPERIMENT_RESULT,
    add_body,
    add_box,
    add_figure,
    add_formula,
    add_list,
    add_page_number,
    add_table,
    add_toc,
    configure_matplotlib,
    connect,
    restart_page_number,
    save_figure,
    set_run_font,
    setup_canvas,
    setup_document,
)
from docx import Document  # noqa: E402
from docx.enum.section import WD_SECTION  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.shared import Pt  # noqa: E402

OUT_DIR = ROOT / "zyjd_system" / "docs" / "fused_report"
DOCX_PATH = OUT_DIR / "智源净域：基于多 Agent 零信任协同与 RAG 知识投毒因果验证的信息污染溯源纠偏系统.docx"


def add_heading(doc: Document, text: str, level: int) -> None:
    """添加与目录层级对应的标题。"""
    doc.add_heading(text, level=level)


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    """批量添加正文段落。"""
    for paragraph in paragraphs:
        add_body(doc, paragraph)


def add_competition_cover(doc: Document) -> None:
    """生成作品赛封面，避免出现学校、院系和指导教师信息。"""
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("第十九届全国大学生信息安全竞赛（作品赛）"), size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("暨第三届“长城杯”网数智安全大赛（作品赛）"), size=16, bold=True)
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("作 品 报 告"), size=26, bold=True)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("□命题赛道             ■自由赛道"), size=12)
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("作品名称"), size=14, bold=True)
    title = "智源净域：基于多 Agent 零信任协同与 RAG 知识投毒因果验证的\n信息污染溯源纠偏系统"
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run(title), size=19, bold=True)
    for _ in range(6):
        doc.add_paragraph()
    for label in ("电子邮箱：____________________________", "提交日期：____________________________"):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        set_run_font(p.add_run(label), size=12)
    doc.add_page_break()


def figure_ipjg() -> Path:
    """绘制 IPJG 的证据、声明、共识、动作四层结构。"""
    fig, ax = setup_canvas((14, 7.4))
    layers = [
        (0.76, "Evidence Layer", ["Document / Source", "Chunk / Evidence", "RAS · GIS · DualRisk"], COLORS["blue"]),
        (0.55, "Claim Layer", ["Atomic Claim", "Evidence Binding", "Claim Provenance DAG"], COLORS["green"]),
        (0.34, "Consensus Layer", ["Agent Vote", "Conflict / Validation", "Evidence-backed BFT"], COLORS["purple"]),
        (0.13, "Action Layer", ["Answer / Advice", "Execution Action", "Rollback State"], COLORS["orange"]),
    ]
    for y, name, modules, color in layers:
        add_box(ax, 0.025, y, 0.16, 0.13, name, color, 12.5)
        for index, module in enumerate(modules):
            x = 0.24 + index * 0.245
            add_box(ax, x, y, 0.19, 0.13, module, "white", 11.5)
            if index < len(modules) - 1:
                connect(ax, (x + 0.19, y + 0.065), (x + 0.245, y + 0.065), color="#666666")
    for x in (0.335, 0.58, 0.825):
        connect(ax, (x, 0.76), (x, 0.68), color="#B22222")
        connect(ax, (x, 0.55), (x, 0.47), color="#B22222")
        connect(ax, (x, 0.34), (x, 0.26), color="#B22222")
    ax.text(0.94, 0.62, "e→c", fontsize=11, color="#8B0000")
    ax.text(0.94, 0.41, "c→con", fontsize=11, color="#8B0000")
    ax.text(0.94, 0.20, "con→a", fontsize=11, color="#8B0000")
    ax.text(0.51, 0.035, "IPJG：证据污染 → 声明污染 → 共识污染 → 动作风险", ha="center", fontsize=14)
    return save_figure(fig, "fig_ipjg_four_layers.png")


def figure_iptcm() -> Path:
    """绘制 IPTCM 检测、溯源、隔离、回滚和重生成闭环。"""
    fig, ax = setup_canvas((15, 7.0))
    labels = [
        ("任务输入与\n多 Agent 规划", COLORS["purple"]),
        ("RAG 检索与\nEvidence 检测", COLORS["blue"]),
        ("Claim 生成与\n零信任封装", COLORS["green"]),
        ("DAG / IPJG\n联合风险分析", COLORS["orange"]),
        ("污染源定位与\nChunk / Agent 隔离", COLORS["red"]),
        ("状态回滚与\n任务重规划", COLORS["purple"]),
        ("可信重检索与\n可信共识", COLORS["blue"]),
        ("可信重生成与\n审计输出", COLORS["yellow"]),
    ]
    positions = [
        (0.03, 0.67), (0.27, 0.67), (0.51, 0.67), (0.75, 0.67),
        (0.75, 0.27), (0.51, 0.27), (0.27, 0.27), (0.03, 0.27),
    ]
    w, h = 0.18, 0.18
    for (text, color), (x, y) in zip(labels, positions):
        add_box(ax, x, y, w, h, text, color, 12)
    for index in range(3):
        x, y = positions[index]
        nx, ny = positions[index + 1]
        connect(ax, (x + w, y + h / 2), (nx, ny + h / 2))
    connect(ax, (0.75 + w / 2, 0.67), (0.75 + w / 2, 0.27 + h))
    for index in range(4, 7):
        x, y = positions[index]
        nx, ny = positions[index + 1]
        connect(ax, (x, y + h / 2), (nx + w, ny + h / 2))
    connect(ax, (0.03 + w / 2, 0.27), (0.03 + w / 2, 0.10), color="#B22222", style="-")
    connect(ax, (0.03 + w / 2, 0.10), (0.96, 0.10), color="#B22222", style="-")
    connect(ax, (0.96, 0.10), (0.96, 0.76), color="#B22222", style="-")
    connect(ax, (0.96, 0.76), (0.93, 0.76), color="#B22222")
    ax.text(0.52, 0.035, "复核不通过时重新进入检测与溯源流程", ha="center", fontsize=12, color="#8B0000")
    return save_figure(fig, "fig_iptcm_closed_loop.png")


def build_report(fig_ipjg: Path, fig_iptcm: Path) -> None:
    """构建严格匹配用户目录的完整作品报告。"""
    experiment = json.loads(EXPERIMENT_RESULT.read_text(encoding="utf-8"))
    rag_result = experiment["rag"]
    agent_rows = experiment["multi_agent"]["rows"]
    doc = Document()
    setup_document(doc)
    add_page_number(doc.sections[0].footer.paragraphs[0])
    add_competition_cover(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(p.add_run("目  录"), size=16, bold=True)
    add_toc(doc)
    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    restart_page_number(body_section, 1)

    add_heading(doc, "摘  要", 1)
    add_paragraphs(
        doc,
        [
            "随着大语言模型应用由单体问答向多 Agent 协同演进，规划、检索、分析、验证与执行被分配给不同角色节点。"
            "局部幻觉、错误推理链、Prompt Infection、恶意指令注入或拜占庭节点行为可能沿通信与任务依赖持续放大，"
            "形成错误共识并驱动高风险动作。与此同时，RAG 知识库已成为多 Agent 获取企业制度、运维文档和安全情报的"
            "关键外部证据源，污染 Document 与 Chunk 可借助检索吸附和答案诱导进入 Agent 推理链，使风险从知识源扩散"
            "至 Claim、共识和动作层。",
            "本作品提出“智源净域”信息污染溯源纠偏系统。系统以 Zero-Trust Claim Envelope 对关键 Claim 进行身份、"
            "角色、权限、证据、父子依赖、时间戳与签名绑定，利用 Claim Provenance DAG 追踪级联传播，并结合 "
            "Propagation Factor、False Consensus Rate、Drift Velocity、Influence Score 与 Byzantine Suspicion "
            "Score 识别异常节点和错误共识。在知识源侧，系统将 RAG 返回内容统一建模为 Evidence 和 Chunk，通过 RAS、"
            "GIS 与 DualRisk 筛选可疑证据，再利用四路反事实验证和 CausalScore 判断其是否真正导致错误。",
            "为贯通知识源与协同决策，本作品构建 Information Pollution Joint Graph（IPJG），统一表达“证据污染—声明"
            "污染—共识污染—动作风险”四层传播链；进一步提出 Information Pollution Tracing and Correction Model"
            "（IPTCM），通过 Evidence-backed BFT Consensus、Chunk 与 Agent 隔离、状态回滚、任务重规划、风险感知"
            "重检索、NLI Claim-Evidence 验证、TrustScore 和可信重生成，实现“检测—溯源—隔离—回滚—重生成”闭环。"
            "原型采用 FastAPI、TF-IDF、启发式 NLI 与 NetworkX 实现，不依赖 GPU 和在线模型接口，适用于安全运营协同、"
            "企业 RAG 知识库治理和大模型应用安全评测。",
        ],
    )
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    set_run_font(
        p.add_run(
            "关键词：多 Agent；零信任协同；级联错误；RAG 知识投毒；因果验证；IPJG；IPTCM；可信纠偏"
        ),
        bold=True,
    )

    add_heading(doc, "第一章 作品概述", 1)
    add_heading(doc, "1.1 背景分析", 2)
    add_paragraphs(
        doc,
        [
            "近年来，大语言模型的能力边界由内容生成和单轮问答快速扩展到工具调用、复杂规划与自主执行。AutoGen、"
            "CrewAI、LangGraph 等多 Agent 框架通过角色分工、共享上下文和循环协作，将复杂任务拆分给规划 Agent、"
            "检索 Agent、分析 Agent、验证 Agent 与执行 Agent。该模式提高了长链任务的自动化程度，也使系统正确性"
            "从单个模型输出问题转变为多节点、多轮次和多依赖关系共同决定的协同安全问题。",
            "多 Agent 系统中的局部错误不会停留在单个节点。上游 Agent 的幻觉、错误事实、错误工具返回或推理偏差，"
            "可能被下游节点当作既定事实继续加工；Prompt Infection 和恶意指令注入还可能借助 Agent 间通信，将不可信"
            "内容嵌入共享上下文。若验证 Agent 与执行 Agent 继续依赖同一错误前提，系统会形成级联错误、错误多数和"
            "伪一致性，最终把错误 Claim 转化为错误处置建议或执行动作。",
            "RAG 已成为多 Agent 获取外部知识的重要方式。企业制度库、运维文档库、安全情报库、产品手册和内部知识"
            "中心为 Agent 提供了超出模型参数知识的 Evidence，使协同系统能够处理时效性强、专业性高的业务任务。"
            "因此，RAG 知识库不再只是问答组件，而是多 Agent 决策链中的关键证据基础设施，其内容完整性、来源独立性"
            "和检索可靠性直接影响上层协同结果。",
            "知识增强也引入了新的信息污染入口。攻击者或异常数据流程可向知识库写入知识投毒内容、间接提示注入文本、"
            "伪造权威结论或高度重复的污染 Chunk；这些内容可能通过关键词覆盖和语义贴合获得异常检索优势，再以答案"
            "诱导进入 Agent 的 Claim。污染 Claim 被其他 Agent 引用、验证或投票后，便形成“证据污染—声明污染—共识"
            "污染”的链式扩散，并可能进一步触发错误动作。",
            "传统单模型幻觉检测通常只判断最终文本是否合理，无法回答错误由哪个 Evidence 引入；最终答案校验无法还原"
            "Claim 在 Agent 间的传播过程；简单多数投票默认各 Agent 的信息来源相互独立，当多个节点引用同一污染 "
            "Chunk 时，反而会把同源错误包装成多数共识；静态关键词或黑名单过滤则难以区分正常热门文档与真正具有"
            "致错因果贡献的投毒证据。",
            "基于上述问题，本作品面向“知识源—智能体—协同决策”全链路构建信息污染溯源纠偏机制。系统先验证 Agent "
            "身份、权限和 Claim 依赖，再检测 RAG Evidence 的吸附性、诱导性和因果贡献，最后通过 IPJG 统一分析污染"
            "如何从 Chunk 传播到 Claim、共识与动作，并利用 IPTCM 完成隔离、回滚、重规划、重检索和可信重生成。"
        ],
    )

    add_heading(doc, "1.2 相关工作", 2)
    related = [
        ("1.2.1 法律背景研究",
         "我国人工智能安全治理已形成以《网络安全法》《数据安全法》《个人信息保护法》《网络数据安全管理条例》"
         "和《生成式人工智能服务管理暂行办法》为重要依据的制度框架，强调网络运行安全、数据处理责任、生成内容"
         "准确可靠和风险可处置。对多 Agent 与 RAG 融合应用而言，制度要求可具体落实为身份与权限可验证、知识来源"
         "可追踪、关键结论有证据、异常决策可阻断以及处置过程可审计。"),
        ("1.2.2 多智能体协同与 Agent 安全研究",
         "现有研究从任务分工、通信协议、拓扑结构和工具调用等角度提升多 Agent 协作能力，也揭示了共享上下文污染、"
         "角色越权、通信篡改和 Prompt Infection 等风险。其启示是安全控制不能只位于输入输出边界，而应进入 Agent "
         "间每一次 Claim 消费过程，采用持续验证和最小权限原则约束消息传播。"),
        ("1.2.3 级联错误、错误共识与拜占庭节点研究",
         "级联错误研究关注局部幻觉如何沿依赖链扩散，拜占庭容错研究关注部分节点任意故障条件下的可靠决策。简单"
         "多数投票只有在投票者相对独立且恶意节点比例受限时才可靠；对于共享同一 RAG Evidence 的 Agent，票数并"
         "不等于证据数量。本作品因此引入 Claim Provenance DAG、传播指标、BSS 和 Evidence-backed BFT Consensus。"),
        ("1.2.4 RAG 知识投毒攻击与防御研究",
         "PoisonedRAG、Pandora、Phantom 等工作表明，少量定向污染文本可能同时影响检索与生成。已有防御包括静态"
         "过滤、鲁棒检索、来源可信度和结果聚合，但单一检测信号容易把高频正常知识误判为污染，也难以证明某个 "
         "Chunk 是否真正改变答案。本作品采用 RAS 与 GIS 双条件筛查，并以反事实验证确认因果影响。"),
        ("1.2.5 反事实因果验证、图谱溯源与可信证据裁决研究",
         "反事实方法通过删除、保留或替换候选因素观察结果变化，适合区分相关性与致错因果；图谱方法能够表达来源、"
         "转载、引用和传播关系；NLI、ALCE、RAGTruth、FActScore 等工作为 Claim-Evidence 支持判断和答案可信评价"
         "提供基础。本作品将这些能力统一到 IPJG，并把因果结果用于隔离与恢复决策。"),
    ]
    for title, text in related:
        add_heading(doc, title, 3)
        add_body(doc, text)

    add_heading(doc, "1.3 作品特色", 2)
    features = [
        ("1.3.1 多 Agent 零信任声明链机制，实现协同决策全过程可追溯与可信验证",
         "系统以 Zero-Trust Claim Envelope 封装关键 Claim，绑定 agent_id、role、permission、evidence_ids、"
         "parent_claim_ids、timestamp、nonce 和 signature；接收方执行“先验证、后使用”，从消息入口阻断身份冒用、"
         "角色越权、重放和无证据高置信结论。"),
        ("1.3.2 信息污染联合溯源图谱，实现错误传播路径分析与污染源精准定位",
         "IPJG 将 Evidence、Chunk、Claim、共识事件与动作放入四层图结构，既能从错误动作反向定位最早污染证据，"
         "也能从高风险 Evidence 正向计算受影响 Claim、Agent 共识和执行状态。"),
        ("1.3.3 RAG 知识投毒双条件检测与反事实因果验证，实现污染证据智能识别",
         "RAS 描述检索吸附，GIS 描述答案诱导，DualRisk 用于候选排序；四路反事实和 CausalScore 进一步验证候选 "
         "Chunk 是否真正导致错误，避免把正常热门内容直接判为投毒。"),
        ("1.3.4 多 Agent 与 RAG 协同纠偏机制，构建可信共识与可信重生成闭环",
         "系统对高风险 Chunk、Claim 和 Agent 分层处置，回滚受污染状态并重新规划任务；可信证据经 NLI 和 "
         "TrustScore 复核后参与 Evidence-backed BFT Consensus 与可信重生成，实现从发现异常到恢复结果的闭环。"),
    ]
    for title, text in features:
        add_heading(doc, title, 3)
        add_body(doc, text)

    add_heading(doc, "1.4 应用前景分析", 2)
    applications = [
        ("1.4.1 面向安全运营多 Agent 协同场景",
         "在告警研判、威胁分析、处置建议和工单执行流程中，系统可追踪每个 Agent 的 Claim、Evidence 与动作依赖，"
         "抑制错误情报或异常节点造成的级联误判，并为高风险动作提供回滚与人工复核依据。"),
        ("1.4.2 面向企业 RAG 知识库可信问答场景",
         "系统可作为企业制度库、运维库、安全知识库的知识源安全检测与可信证据恢复模块，对污染 Chunk、同源转载"
         "和错误引用进行检测、隔离和重检索，提高知识增强应用的可解释性与可审计性。"),
        ("1.4.3 面向大模型应用可信治理与安全评测场景",
         "IPJG 与 IPTCM 可用于构建多 Agent + RAG 的复合安全评测基线，覆盖错误传播、伪多数共识、知识投毒因果、"
         "源头定位和纠偏恢复等指标，为模型应用上线前测试和运行期治理提供统一分析框架。"),
    ]
    for title, text in applications:
        add_heading(doc, title, 3)
        add_body(doc, text)

    add_heading(doc, "第二章 作品设计与实现", 1)
    add_heading(doc, "2.1 系统方案", 2)
    add_body(
        doc,
        "系统按照“先多 Agent 级联错误治理、再 RAG 知识源安全检测、最后联合溯源纠偏”的顺序组织。上层模块"
        "负责 Claim 可信传递、传播分析和共识恢复；下层模块作为知识源安全检测与可信证据恢复能力，为 Agent 提供"
        "经过风险验证的 Evidence；IPJG 与 IPTCM 负责跨层关联和闭环处置。"
    )

    add_heading(doc, "2.1.1 多 Agent 级联错误检测与可信纠偏模块设计", 3)
    sections_211 = [
        ("2.1.1.1 零信任声明包生成与校验",
         "关键自然语言消息先拆分为原子 Claim，再封装为 Zero-Trust Claim Envelope。系统依次校验 Agent 身份、"
         "角色权限、签名完整性、nonce 防重放、Evidence 存在性和父 Claim 闭合性；任一硬约束失败时，声明不得进入"
         "共享状态。"),
        ("2.1.1.2 Claim Provenance DAG 声明传播图谱构建",
         "系统以 Claim 为节点，以 derives、supports、contradicts、references 和 triggers 为有向边，记录生成者、"
         "接收者、时间、任务轮次与证据绑定。该 DAG 支持从异常 Claim 向后计算影响范围，也支持从错误结果向前回溯"
         "最早异常声明。"),
        ("2.1.1.3 级联错误检测与拜占庭节点识别",
         "系统联合 Propagation Factor、False Consensus Rate、Drift Velocity、Influence Score、权限异常和证据"
         "矛盾率识别异常传播；对持续制造错误、异常放大或绕过验证的 Agent 计算 Byzantine Suspicion Score，区分"
         "普通失误、共享证据污染与主动异常行为。"),
        ("2.1.1.4 证据加权可信共识机制",
         "Evidence-backed BFT Consensus 不直接按 Agent 数量计票，而是按 Evidence 的 TrustScore、Claim-Evidence "
         "支持度、来源独立性和 Agent 的 BSS 赋权。同一 Document 的复制内容会进行来源聚类，避免重复转述形成伪多数。"),
        ("2.1.1.5 隔离、回滚与任务重规划纠偏机制",
         "高风险 Claim 被撤销或降权，可疑 Agent 被限制角色或隔离；若污染已写入共享状态，系统回滚至最近可信 "
         "Checkpoint，清理受影响后继状态，并将任务重新分配给剩余可信 Agent，形成可审计恢复记录。"),
    ]
    for title, text in sections_211:
        add_heading(doc, title, 4)
        add_body(doc, text)

    add_heading(doc, "2.1.2 RAG 知识投毒检测、因果溯源与可信重生成模块设计", 3)
    sections_212 = [
        ("2.1.2.1 RAG 证据对象建模与知识库接入",
         "企业知识库、运维文档和安全情报首先按 Document 与 Chunk 管理，再统一映射为 Evidence。Evidence 记录"
         "内容、来源、时间、哈希、检索历史、标签、Claim 绑定关系以及 RAS、GIS、DualRisk、CausalScore 等评分。"),
        ("2.1.2.2 检索吸附性与答案诱导性检测",
         "RAS 衡量 Chunk 相对随机基线的检索频率，GIS 衡量答案对该 Chunk 的相对文本依赖，DualRisk 联合两个"
         "必要条件筛选重点候选。该模块作为 Agent 使用知识前的安全门，不作为独立问答系统。"),
        ("2.1.2.3 四路反事实因果验证",
         "系统固定查询与生成策略，分别构造原始 Top-K、删除可疑、仅可疑和可信替代四路 Evidence 上下文。通过"
         "答案差异、正确性变化和 CausalScore 判断候选 Chunk 是否具有实际致错贡献。"),
        ("2.1.2.4 RAG 投毒传播图谱构建",
         "知识侧图谱记录 Page、Document、Chunk、Query、Claim 和 Answer 之间的 contains、retrieved_by、supports、"
         "contradicts、similar_to、copied_from、same_claim 与 caused_error 等关系，并作为 IPJG Evidence Layer 的输入。"),
        ("2.1.2.5 TrustScore 可信评分与可信重生成",
         "高风险 Chunk 及其相似副本被隔离后，系统执行风险感知重检索，优先选择来源独立、时间有效且与 Claim 形成"
         "支持关系的 Evidence。可信重生成只使用通过 NLI 与 TrustScore 门限的证据，并对证据不足内容明确标注不确定性。"),
    ]
    for title, text in sections_212:
        add_heading(doc, title, 4)
        add_body(doc, text)

    add_heading(doc, "2.2 实现原理", 2)
    add_heading(doc, "2.2.1 多 Agent 零信任协同机制", 3)
    sections_221 = [
        ("2.2.1.1 Agent 身份、角色与权限绑定",
         "每个 Agent 具有唯一身份、职责集合和最小权限策略。规划 Agent 可拆分任务但不能直接执行高风险动作，验证 "
         "Agent 可给出证据裁决但不能修改原始证据，执行 Agent 只能消费通过共识门限的 Claim。"),
        ("2.2.1.2 原子声明 Claim 抽取与证据绑定",
         "系统将复合结论拆为可独立验证的主语—谓词—宾语 Claim，并绑定 Evidence ID、父 Claim、生成 Agent 和"
         "置信信息。缺少 Evidence 的事实性 Claim 被标记为待补证，不得仅凭模型自报置信度进入可信共识。"),
        ("2.2.1.3 声明传播路径追踪与风险评分",
         "Claim Provenance DAG 记录每次引用、派生、冲突和验证。节点风险由内容矛盾、父节点风险、传播深度、漂移"
         "速度和生成 Agent 的 BSS 共同决定，并在新增 Evidence 或风险变化时增量更新。"),
    ]
    for title, text in sections_221:
        add_heading(doc, title, 4)
        add_body(doc, text)

    add_heading(doc, "2.2.2 RAG 知识投毒检测机制", 3)
    add_heading(doc, "2.2.2.1 RAS 检索吸附性评分", 4)
    add_body(doc, "RAS 用于衡量某一 Chunk 的历史检索频率相对随机检索基线的放大程度：")
    add_formula(doc, "RAS(chunk_i) = (freq_i / total_retrievals) / (1 / total_chunks)")
    add_body(doc, "其中 freq_i 为 chunk_i 被检索次数，total_retrievals 为所有 Chunk 检索总次数，total_chunks 为语料库 Chunk 数。RAS > 1 表示其检索频率高于随机基线。")
    add_heading(doc, "2.2.2.2 GIS 答案诱导性评分", 4)
    add_formula(doc, "GIS(chunk_i, answer) = sim(answer, chunk_i) / max_sim(answer, all_chunks)")
    add_body(doc, "sim() 在 MVP 中使用 TF-IDF 余弦相似度。GIS 接近 1 表示答案在当前候选集合中高度贴近该 Chunk。")
    add_heading(doc, "2.2.2.3 DualRisk 双条件投毒风险评分", 4)
    add_formula(doc, "DualRisk(chunk_i) = RAS(chunk_i) × GIS(chunk_i)")
    add_body(doc, "实现中先将 RAS 归一化到 [0,1] 再与 GIS 组合；只有检索异常与生成诱导同时显著时，Chunk 才进入反事实验证。")
    add_heading(doc, "2.2.2.4 CausalScore 因果影响评分", 4)
    add_formula(doc, "CausalScore(chunk_i) = 1 - sim(A_orig, A_remove) / sim(A_orig, A_only_suspect)")
    add_body(doc, "A_orig、A_remove 和 A_only_suspect 分别表示原始、删除可疑和仅可疑 Evidence 生成答案。实现中使用 ε 保护分母并将结果截断到 [0,1]。")

    add_heading(doc, "2.3 基于零信任声明链与因果证据验证的信息污染联合溯源纠偏模型实现", 2)
    add_body(
        doc,
        "本节提出 IPJG 与 IPTCM，将多 Agent 级联错误和 RAG 知识投毒统一描述为跨层信息污染传播问题。"
        "IPJG 负责保存“污染从何处进入、经过哪些 Claim 和共识事件、影响哪些动作”的结构证据；IPTCM 负责根据"
        "结构证据执行动态复检、根因定位、隔离、回滚和可信恢复。"
    )
    add_heading(doc, "2.3.1 基于证据—声明关联关系的污染传播建模", 3)
    add_paragraphs(
        doc,
        [
            "设一次任务中参与判断的 Evidence 集合和 Claim 集合分别为：",
        ],
    )
    add_formula(doc, "E = {e_1, e_2, ..., e_n}")
    add_formula(doc, "C = {c_1, c_2, ..., c_m}")
    add_body(doc, "每个 Claim 必须记录其依赖 Evidence 子集。例如：")
    add_formula(doc, "Dep(c_j) = {e_1, e_4, e_7, ...}")
    add_body(doc, "系统通过 NLI 或启发式规则计算 Evidence 对 Claim 的支持概率：")
    add_formula(doc, "Support(c, e) = P(entailment | c, e)")
    add_paragraphs(
        doc,
        [
            "Evidence 与 Claim 由依赖和裁决关系组成二部图。Evidence 节点保存 source_id、document_id、chunk_id、"
            "内容哈希、检索排名和风险分数；Claim 节点保存生成 Agent、父 Claim、任务轮次和当前风险。边保存 supports、"
            "contradicts、neutral、引用时间和绑定强度。",
            "当 Support(c,e) 低于阈值且 Evidence 具有较高 DualRisk 时，系统建立污染传播边 e_i → c_j。该边不是"
            "仅表示文本不一致，而是表示高风险 Evidence 已被 Agent 用于形成 Claim；若四路反事实进一步得到较高 "
            "CausalScore，则边属性 causal_confirmed 被置为真，作为后续隔离和回滚的重要依据。",
            "通过该二部图，知识污染向声明污染的过程被显式化：同一污染 Chunk 可连接多个 Claim，多个转载 Evidence "
            "也可聚合到同一错误 Claim。系统据此区分“多个独立证据支持同一 Claim”和“多个 Agent 重复引用同一污染"
            "来源”两种表面相似、实质不同的共识结构。"
        ],
    )

    add_heading(doc, "2.3.2 基于级联传播行为的动态复检触发建模", 3)
    add_body(doc, "为使 Agent 级联异常能够反向触发知识源复检，系统定义 Claim 的传播因子：")
    add_formula(doc, "PF(c_i) = N_prop(c_i) / N_total")
    add_body(doc, "N_prop(c_i) 表示直接或间接依赖 c_i 的后继 Claim 数，N_total 表示当前任务全部 Claim 数。PF 越高，说明该 Claim 对协同链影响范围越大。")
    add_body(doc, "系统同时计算 Claim 风险随时间或协同轮次的漂移速度：")
    add_formula(doc, "DV(c_i) = ΔRisk / Δt")
    add_paragraphs(
        doc,
        [
            "当 PF(c_i) > τ_1 且 DV(c_i) > τ_2 时，系统认为该 Claim 正在快速扩散并产生风险漂移，自动触发 RAG "
            "证据复检流程。复检对象不是整个知识库，而是 Dep(c_i) 及其相似 Chunk、copied_from 来源和 same_claim "
            "关联 Evidence，从而控制计算开销。",
            "复检依次重新执行 RAS、GIS、DualRisk 和四路反事实因果验证。若 Evidence 风险上升，则沿 Evidence→Claim "
            "边更新 Claim 风险并冻结相关动作；若 Evidence 仍可信，则系统转而检查生成 Agent 的推理错误、身份权限"
            "和通信路径。该双向触发机制避免把所有级联错误都归因于知识库，也避免知识风险只在初次检索时检测一次。"
        ],
    )

    add_heading(doc, "2.3.3 基于联合传播图谱的信息污染溯源核心结构设计", 3)
    add_body(doc, "IPJG 采用有向异构属性图表示四层污染传播结构：")
    add_formula(doc, "G = (V, E)")
    add_formula(doc, "V = {V_E, V_C, V_Con, V_A}")
    add_formula(doc, "E = {E_e→c, E_c→c, E_c→con, E_con→a}")
    add_table(
        doc, 1, "IPJG 四层节点与关键属性",
        ["层级", "核心节点", "关键属性", "安全语义"],
        [
            ["Evidence Layer", "Source、Document、Chunk、Evidence", "来源、哈希、时间、RAS/GIS/DualRisk/CausalScore", "描述知识入口与证据风险"],
            ["Claim Layer", "Claim、Agent、父子依赖", "生成者、Evidence 绑定、PF、DV、风险", "描述声明形成与级联传播"],
            ["Consensus Layer", "Vote、Validation、Conflict、Consensus", "票权、来源簇、BSS、共识状态", "描述错误共识与伪多数"],
            ["Action Layer", "Answer、Advice、Action、Checkpoint", "执行状态、风险等级、回滚点", "描述污染造成的业务影响"],
        ],
    )
    add_figure(doc, fig_ipjg, 1, "IPJG 信息污染联合溯源图谱四层结构")
    add_body(doc, "为确定优先复核和隔离对象，系统定义节点综合影响力：")
    add_formula(doc, "IS(v) = α·Risk(v) + β·Centrality(v) + γ·Propagation(v)")
    add_body(doc, "Risk(v) 表示节点自身风险，Centrality(v) 表示其在 IPJG 中的结构中心性，Propagation(v) 表示其影响后继节点的能力；α、β、γ 根据场景配置且和为 1。")
    add_body(doc, "对 Agent a 定义拜占庭嫌疑评分：")
    add_formula(doc, "BSS(a) = λ_1·ErrRate + λ_2·ConflictRate + λ_3·AbnormalPropagation")
    add_body(doc, "ErrRate、ConflictRate 和 AbnormalPropagation 分别表示历史错误率、与可信证据冲突率和异常传播率。BSS 用于共识降权和隔离优先级，不以单次分歧直接认定恶意。")
    add_body(doc, "Evidence-backed BFT Consensus 对 Claim 的有效权重定义为：")
    add_formula(doc, "Weight(c) = TrustScore(e) × Support(c,e) × (1 - BSS(a))")
    add_paragraphs(
        doc,
        [
            "该机制优于简单多数投票的关键在于投票权来源于证据而非节点数量。若多个 Agent 引用同一污染 Document "
            "或其复制 Chunk，系统先通过 copied_from、内容哈希和 same_claim 进行来源聚类，再按一个证据簇计算"
            "独立性权重；因此，同源错误不会因被多个 Agent 转述而线性增加票权。",
            "对于少数但持有高 TrustScore 独立 Evidence 的 Agent，其 Claim 可获得高于多数低质量同源证据的权重。"
            "当共识仍不足时，系统输出“证据不足”并触发补证，而不是强制选择票数较多的结论。由此能够抑制共享污染"
            "证据造成的 False Consensus，同时保留对正常少数意见的验证空间。"
        ],
    )

    add_heading(doc, "2.3.4 本文提出的信息污染联合溯源纠偏模型", 3)
    add_paragraphs(
        doc,
        [
            "IPTCM 以 IPJG 为状态载体，以零信任校验、知识投毒因果验证和可信恢复为三个核心控制环。完整流程为："
            "用户任务输入 → 多 Agent 协同规划 → RAG 检索知识库 → Evidence 风险检测 → Agent 生成 Claim → "
            "零信任声明链封装 → Claim Provenance DAG 传播追踪 → 联合图谱风险分析 → 污染源定位 → 隔离污染 "
            "Chunk 或可疑 Agent → 状态回滚与任务重规划 → 可信证据重检索 → 可信共识形成 → 可信结果输出。",
            "检测阶段同时观察 Agent 主体可信、Claim 依赖风险和 Evidence 风险；溯源阶段从高风险 Answer 或 Action "
            "沿 E_con→a、E_c→con、E_c→c 和 E_e→c 反向遍历，得到最小污染解释子图，并按 IS 排序根因节点；隔离"
            "阶段根据 causal_confirmed、BSS 和动作影响范围选择 Chunk、Claim 或 Agent 粒度，避免过度处置。",
            "回滚阶段恢复最近一个所有关键 Claim 均通过零信任校验且 Evidence 未被隔离的 Checkpoint，重新规划时"
            "禁止复用受污染上下文；可信重检索阶段扩大候选池并加入风险惩罚、来源多样性和时间有效性约束；可信共识"
            "阶段使用 Evidence-backed BFT Consensus；最终答案再经 NLI Claim-Evidence 验证和 TrustScore 复核。",
            "因此，IPTCM 不是告警型过滤器，而是实现“检测—溯源—隔离—回滚—重生成”的闭环治理体系。每次处置均"
            "保留触发条件、根因路径、隔离对象、回滚点、重检索 Evidence 和最终 Claim，便于比赛演示、人工复核和"
            "责任审计。"
        ],
    )
    add_figure(doc, fig_iptcm, 2, "IPTCM 信息污染联合溯源纠偏闭环")

    add_heading(doc, "2.4 系统实现架构", 2)
    architecture = [
        ("2.4.1 Agent 编排与任务调度模块",
         "负责角色注册、任务分解、消息路由、轮次控制和 Checkpoint 管理，并向安全模块提交 Claim、Evidence 引用和动作请求。"),
        ("2.4.2 零信任通信与声明链校验模块",
         "负责 Zero-Trust Claim Envelope 生成、签名与防重放校验、角色权限判断、父 Claim 闭合检查和 Provenance 事件记录。"),
        ("2.4.3 RAG 检索与知识库管理模块",
         "负责 Document 接入、Chunk 切分去重、检索历史、Evidence 对象生成、风险感知重排序和隔离列表执行。"),
        ("2.4.4 图谱存储与传播路径分析模块",
         "MVP 使用 NetworkX 保存 Claim Provenance DAG 与 IPJG，提供前向影响分析、反向根因搜索、来源聚类和关键路径排序。"),
        ("2.4.5 风险评分、可信共识、纠偏决策与可视化审计报告模块",
         "统一计算 RAS、GIS、DualRisk、CausalScore、PF、DV、IS、BSS 与 TrustScore，执行 Evidence-backed BFT "
         "Consensus 和纠偏策略，并输出图谱、评分、处置日志与结构化风险报告。"),
    ]
    for title, text in architecture:
        add_heading(doc, title, 3)
        add_body(doc, text)

    add_heading(doc, "2.5 核心算法流程", 2)
    algorithms = [
        ("2.5.1 多 Agent 级联错误检测算法",
         ["校验 Zero-Trust Claim Envelope；", "更新 Claim Provenance DAG；", "计算 PF、DV、冲突率与父节点风险；",
          "更新 Agent 的 BSS；", "输出高风险 Claim、可疑 Agent 和受影响动作集合。"]),
        ("2.5.2 RAG 知识投毒因果验证算法",
         ["构建 Evidence 与 Chunk；", "根据检索历史计算 RAS；", "根据答案相似度计算 GIS；", "计算 DualRisk 并筛选候选；",
          "执行四路反事实并计算 CausalScore；", "输出因果确认的污染 Chunk。"]),
        ("2.5.3 信息污染联合溯源纠偏算法",
         ["将 Evidence、Claim、Consensus 和 Action 写入 IPJG；", "按 IS 搜索高影响根因；", "确定最小隔离集合；",
          "回滚可信 Checkpoint 并重规划；", "风险感知重检索；", "Evidence-backed BFT Consensus、可信重生成与审计输出。"]),
    ]
    for title, steps in algorithms:
        add_heading(doc, title, 3)
        add_list(doc, steps, numbered=True)

    add_heading(doc, "2.6 性能指标", 2)
    add_table(
        doc, 2, "系统性能与安全指标",
        ["维度", "指标", "计算口径", "当前状态"],
        [
            ["级联检测", "Propagation Factor / Drift Velocity", "Claim 影响范围与风险变化速率", "原型已实现思路"],
            ["错误共识", "False Consensus Rate", "错误 Claim 进入最终共识的比例", "公开集待实测"],
            ["节点识别", "Byzantine Detection F1", "可疑 Agent 识别精确率、召回率与 F1", "本地仿真实测"],
            ["投毒检测", "Poison Detection F1", "投毒 Chunk 识别质量", "本地案例实测"],
            ["因果溯源", "Source Localization Rate", "污染 Evidence / Agent 根因定位率", "本地仿真实测"],
            ["纠偏恢复", "Correction Success Rate", "错误发生后恢复正确任务结果的比例", "本地仿真实测"],
            ["可信质量", "TrustScore / Evidence Support Rate", "重生成答案的证据支持与综合可信度", "本地案例实测"],
            ["系统开销", "Latency / Memory / Graph Size", "端到端时延、内存和图规模", "待实测"],
        ],
    )

    add_heading(doc, "第三章 作品测试与分析", 1)
    add_heading(doc, "3.1 测试环境", 2)
    add_table(
        doc, 3, "原型系统测试环境",
        ["类别", "配置"],
        [
            ["系统环境", "Windows 11 / Linux，Python 3.10+"],
            ["计算资源", "普通 CPU 环境，不要求 GPU"],
            ["核心依赖", "FastAPI、Pydantic v2、NetworkX、scikit-learn、NumPy、pytest"],
            ["模型策略", "TF-IDF 相似度、启发式 NLI、模板化可信重生成"],
            ["网络边界", "不调用在线模型 API，不连接真实业务服务"],
        ],
    )

    add_heading(doc, "3.2 测试数据集", 2)
    add_heading(doc, "3.2.1 公开数据集", 3)
    datasets = [
        ("3.2.1.1 MAST-Data、MultiAgentBench 与 AgentDojo 数据集",
         "MAST-Data 用于多 Agent 失败轨迹、错误类型和传播根因分析；MultiAgentBench 用于评估协作任务完成、"
         "角色分工和通信过程；AgentDojo 用于构造不可信工具返回、间接提示注入和权限约束场景。三者分别支撑"
         "Claim 级错误定位、协同基线对比和零信任通信测试。"),
        ("3.2.1.2 PoisonedRAG 与 RAGTruth 数据集",
         "PoisonedRAG 用于构造定向知识投毒、不同投毒比例和 Top-K 条件下的检测与因果验证；RAGTruth 提供 RAG "
         "输出幻觉与不忠实内容标注，用于评估 NLI Claim-Evidence 验证和可信重生成质量。"),
        ("3.2.1.3 ALCE 与 CFEVER 数据集",
         "ALCE 用于评估回答引用完整性、引用正确性和可验证性；CFEVER 用于中文 Claim 的支持、反驳和证据不足"
         "三分类。二者作为纠偏后答案质量评测集，不作为知识投毒攻击样本。"),
    ]
    for title, text in datasets:
        add_heading(doc, title, 4)
        add_body(doc, text)

    add_heading(doc, "3.2.2 测试场景构建", 3)
    scenarios = [
        ("3.2.2.1 基于公开数据集的多 Agent 协同任务场景",
         "将公开任务拆分为规划、检索、分析、验证和执行角色，保留每轮 Claim、Evidence 和消息依赖，构造正常协作、"
         "单点错误、错误多数与异常节点比例变化场景。"),
        ("3.2.2.2 基于 PoisonedRAG 的知识投毒场景",
         "按原数据设置离线构造污染 Document 与 Chunk，测试 RAS、GIS、DualRisk、四路反事实、CausalScore 和"
         "污染源定位；公开数据集结果在完成正式复现前统一标记为待实测。"),
        ("3.2.2.3 基于 RAGTruth 的幻觉与错误引用场景",
         "把响应拆分为 Claim，与上下文 Evidence 建立支持、矛盾和中立关系，评价 Claim-Evidence 矩阵、TrustScore "
         "与可信重生成是否减少无证据 Claim。"),
        ("3.2.2.4 基于 AgentDojo 的通信干扰与 Prompt Infection 场景",
         "将不可信工具返回作为外部 Evidence 输入，观察间接提示是否跨 Agent 传播，验证角色权限、声明封装、动作门控"
         "和回滚机制。"),
        ("3.2.2.5 基于 CFEVER 的困难负样本与事实冲突场景",
         "选取语义相近但标签不同的中文样本，测试 NLI 启发式在否定、时间冲突、实体替换和证据不足条件下的误报与漏报。"),
    ]
    for title, text in scenarios:
        add_heading(doc, title, 4)
        add_body(doc, text)

    add_heading(doc, "3.3 测试方案", 2)
    schemes = [
        ("3.3.1 多 Agent 模块测试方案",
         ["正常协作与正常分歧测试；", "单 Agent 幻觉和级联传播测试；", "拜占庭节点比例变化测试；",
          "共享污染 Evidence 形成伪多数测试；", "身份、权限、签名、防重放和 Checkpoint 回滚测试。"]),
        ("3.3.2 RAG 知识投毒模块测试方案",
         ["不同 Top-K、投毒比例和检索历史下的 RAS 测试；", "不同答案依赖程度下的 GIS 测试；",
          "DualRisk 阈值和权重消融；", "四路反事实与真实投毒标签一致性测试；", "Chunk 隔离、相似副本排除和可信重生成测试。"]),
        ("3.3.3 多 Agent 与 RAG 融合测试方案",
         ["污染 Chunk → Claim → 共识 → 动作端到端传播测试；", "高 PF 与高 DV 触发 Evidence 动态复检测试；",
          "IPJG 反向根因定位测试；", "Evidence-backed BFT Consensus 与多数投票对比；",
          "隔离、回滚、重规划和重生成闭环恢复测试。"]),
        ("3.3.4 评估指标",
         ["任务正确率、False Consensus Rate、源头定位率和纠偏成功率；", "投毒检测 Precision、Recall、F1 和 FPR；",
          "拜占庭节点识别 Precision、Recall 与 F1；", "TrustScore、Evidence Support Rate 和引用正确性；",
          "端到端时延、图谱规模和处置开销。"]),
    ]
    for title, items in schemes:
        add_heading(doc, title, 3)
        add_list(doc, items, numbered=True)

    add_heading(doc, "3.4 测试结果", 2)
    best_rag = rag_result["best_baselines"]
    selected_agents = [
        row for row in agent_rows
        if row["risk_threshold"] == 0.55 and row["malicious_ratio"] in (0.14, 0.29, 0.43)
    ]
    add_heading(doc, "3.4.1 多 Agent 级联错误检测结果", 3)
    add_table(
        doc, 4, "原型系统本地受控多 Agent 实验结果",
        ["异常 Agent 比例", "无防护", "多数投票", "单 Verifier", "本文方法", "纠偏成功率"],
        [[f'{r["malicious_ratio"]:.2f}', f'{r["no_defense_accuracy"]:.3f}', f'{r["majority_accuracy"]:.3f}',
          f'{r["single_verifier_accuracy"]:.3f}', f'{r["proposed_accuracy"]:.3f}', f'{r["correction_rate"]:.3f}']
         for r in selected_agents],
    )
    add_body(doc, "该结果来自 7 个 Agent、固定随机种子 20260614、每组 500 次的规则受控仿真，只用于验证原型流程可运行，不代表公开数据集性能。")
    add_figure(doc, EXPERIMENT_FIGURE_DIR / "exp_multi_agent_robustness.png", 3, "不同异常 Agent 比例下各方法任务正确率")
    add_heading(doc, "3.4.2 拜占庭节点识别结果", 3)
    add_body(doc, "当前原型已记录 source_localization_rate 作为异常源命中率，并完成风险阈值扫描；公开数据集上的拜占庭节点 Precision、Recall 与 F1 尚未完成，表中统一标记为待实测。")
    add_table(
        doc, 5, "拜占庭节点识别结果状态",
        ["数据来源", "Precision", "Recall", "F1", "说明"],
        [["本地受控仿真", "待补充分项统计", "待补充分项统计", "待补充分项统计", "已保存源头命中率与原始 CSV"],
         ["MAST-Data / MultiAgentBench", "待实测", "待实测", "待实测", "不得以本地仿真替代公开集结果"]],
    )
    add_figure(doc, EXPERIMENT_FIGURE_DIR / "exp_multi_agent_threshold.png", 4, "联合风险阈值对任务正确率与源头定位率的影响")
    add_heading(doc, "3.4.3 RAG 知识投毒检测结果", 3)
    add_table(
        doc, 6, "原型系统本地受控 RAG 投毒检测结果",
        ["方法", "最优阈值", "Precision", "Recall", "F1", "FPR"],
        [[name, f'{r["threshold"]:.2f}', f'{r["precision"]:.3f}', f'{r["recall"]:.3f}', f'{r["f1"]:.3f}', f'{r["fpr"]:.3f}']
         for name, r in best_rag.items()],
    )
    add_body(doc, f"本地实验包含 {rag_result['sample_count']} 条 Top-K Evidence，其中投毒 Evidence {rag_result['poison_count']} 条。样本规模较小，结果仅说明原型可运行；PoisonedRAG 公开数据集结果待实测。")
    add_figure(doc, EXPERIMENT_FIGURE_DIR / "exp_dualrisk_threshold.png", 5, "DualRisk 阈值对 Precision、Recall 与 F1 的影响")
    add_heading(doc, "3.4.4 因果溯源与可信重生成结果", 3)
    add_table(
        doc, 7, "本地受控案例 TrustScore 纠偏前后对比",
        ["案例", "纠偏前", "纠偏后", "提升", "隔离投毒 Evidence"],
        [[("企业制度投毒" if r["case_id"] == "case_enterprise_rag_poisoning" else "站群投毒"),
          f'{r["trust_before"]:.2f}', f'{r["trust_after"]:.2f}', f'{r["trust_gain"]:.2f}',
          f'{r["isolated_count"]}/{r["poison_count"]}'] for r in rag_result["case_results"]],
    )
    add_figure(doc, EXPERIMENT_FIGURE_DIR / "exp_trust_before_after.png", 6, "风险隔离与可信重生成前后的 TrustScore")
    add_heading(doc, "3.4.5 融合纠偏前后对比结果", 3)
    add_table(
        doc, 8, "融合纠偏结果与公开数据集实测状态",
        ["评价项", "纠偏前", "纠偏后", "数据状态"],
        [
            ["多 Agent 任务正确率", "无防护/多数投票基线", "Evidence-backed BFT Consensus", "本地受控仿真实测"],
            ["RAG TrustScore", "保留高风险 Evidence", "隔离后可信重生成", "本地受控案例实测"],
            ["False Consensus Rate", "待实测", "待实测", "公开数据集待实测"],
            ["端到端 IPTCM 恢复率", "待实测", "待实测", "复合公开场景待实测"],
        ],
    )

    add_heading(doc, "3.5 模型与方法对比", 2)
    comparisons = [
        ("3.5.1 与无防护多 Agent 系统对比", "无防护系统直接消费上游消息和 RAG Evidence，无法阻断错误 Claim。本文方法增加声明校验、传播检测和恢复控制。"),
        ("3.5.2 与多数投票方法对比", "多数投票只统计节点数量，无法识别同源污染。Evidence-backed BFT Consensus 按证据可信、支持关系、来源独立性和 BSS 赋权。"),
        ("3.5.3 与单 Verifier Agent 方法对比", "单 Verifier 存在单点失效且难以覆盖长链传播。本文通过 Claim DAG、IPJG 和多证据裁决提供结构化验证。"),
        ("3.5.4 与静态 RAG 过滤方法对比", "静态过滤只能发现表面异常，无法证明致错。本文联合 RAS、GIS、DualRisk 与四路反事实，并将结果用于回滚和重生成。"),
    ]
    for title, text in comparisons:
        add_heading(doc, title, 3)
        add_body(doc, text)

    add_heading(doc, "3.6 测试结果分析", 2)
    add_paragraphs(
        doc,
        [
            "本地受控结果表明，证据加权共识在异常 Agent 比例提高时仍保持较高任务正确率，说明来源独立性和风险"
            "门控能够抑制部分伪多数；DualRisk 在小规模案例中实现了较高召回率，说明双条件筛查可覆盖已构造的污染"
            "Evidence；隔离与可信重生成后 TrustScore 上升，验证了知识源恢复流程的可运行性。",
            "上述结论存在明确边界：多 Agent 数据来自固定规则仿真，RAG 样本数量较少，启发式 NLI 对复杂语义冲突的"
            "适应性有限，公开数据集与端到端复合场景尚未完成正式复现。因此报告不将本地数值外推为通用性能，所有"
            "MAST-Data、MultiAgentBench、AgentDojo、PoisonedRAG、RAGTruth、ALCE 和 CFEVER 结果均保留待实测状态。"
        ],
    )

    add_heading(doc, "第四章 创新性说明", 1)
    innovations = [
        ("4.1 面向多 Agent 协同的零信任声明链机制",
         "将零信任从网络访问扩展到 Claim 消费，使用身份、角色、权限、Evidence、父子依赖和签名共同约束协同消息。"),
        ("4.2 面向级联错误的 Claim 传播图谱溯源方法",
         "以 Claim Provenance DAG 表达多轮推理依赖，并用 PF、DV 和 IS 识别快速扩散、高影响的错误声明。"),
        ("4.3 面向 RAG 知识投毒的双条件检测与反事实因果验证",
         "通过 RAS 与 GIS 联合筛查检索和生成两个必要条件，再通过四路反事实与 CausalScore 确认真正致错的 Chunk。"),
        ("4.4 面向错误共识的证据加权可信纠偏机制",
         "Evidence-backed BFT Consensus 将 TrustScore、Support、来源独立性与 BSS 纳入票权，抑制共享污染证据造成的伪多数。"),
        ("4.5 多 Agent 与 RAG 融合的信息污染溯源纠偏闭环",
         "IPJG 与 IPTCM 贯通 Evidence、Claim、Consensus 和 Action，实现从知识源定位到 Agent 隔离、状态回滚、可信重检索和可信重生成的统一治理。"),
    ]
    for title, text in innovations:
        add_heading(doc, title, 2)
        add_body(doc, text)

    add_heading(doc, "第五章 总结", 1)
    add_heading(doc, "5.1 工作总结", 2)
    add_body(
        doc,
        "本作品构建了面向“多 Agent + RAG”融合场景的信息污染溯源纠偏系统。它不是普通 RAG 问答系统，也不是"
        "普通多 Agent 编排系统，而是围绕证据污染、声明污染、共识污染和动作风险建立零信任验证、因果检测、联合"
        "溯源与可信恢复机制。系统通过 Zero-Trust Claim Envelope、Claim Provenance DAG、IPJG、IPTCM、"
        "Evidence-backed BFT Consensus 和可信重生成，形成可运行、可展示、可解释的安全闭环。"
    )
    add_heading(doc, "5.2 未来展望", 2)
    add_body(
        doc,
        "后续工作将优先完成公开数据集可复现实测，补充 False Consensus Rate、拜占庭节点 F1、端到端时延和"
        "大规模图谱性能；在保持接口稳定的前提下，可替换更强的向量检索与 NLI 模型，并完善企业身份系统、知识库"
        "版本管理、人工复核和审计导出能力。扩展仍将限定在多 Agent 与 RAG 信息污染治理方向。"
    )

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] NIST. Zero Trust Architecture: NIST Special Publication 800-207[R]. 2020.",
        "[2] LEWIS P, PEREZ E, PIKTUS A, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks[C]//NeurIPS. 2020.",
        "[3] ZOU W, GENG R, WANG B, JIA J. PoisonedRAG: Knowledge Corruption Attacks to Retrieval-Augmented Generation of Large Language Models[C]//USENIX Security. 2025.",
        "[4] DENG G, LIU Y, WANG K, et al. Pandora: Jailbreak GPTs by Retrieval Augmented Generation Poisoning[J]. arXiv:2402.08416, 2024.",
        "[5] CHAUDHARI H, SEVERI G, ABASCAL J, et al. Phantom: General Trigger Attacks on Retrieval Augmented Language Generation[J]. arXiv:2405.20485, 2024.",
        "[6] ZHANG B, XIN H, FANG M, et al. Traceback of Poisoning Attacks to Retrieval-Augmented Generation[C]//The Web Conference. 2025.",
        "[7] XIANG C, WU T, ZHONG Z, et al. Certifiably Robust RAG against Retrieval Corruption[J]. arXiv:2405.15556, 2024.",
        "[8] GAO T, YEN H, YU J, CHEN D. Enabling Large Language Models to Generate Text with Citations[C]//EMNLP. 2023.",
        "[9] NIU C, WU Y, ZHU J, et al. RAGTruth: A Hallucination Corpus for Developing Trustworthy Retrieval-Augmented Language Models[J]. arXiv:2401.00396, 2024.",
        "[10] MIN S, KRISHNA K, LYU X, et al. FActScore: Fine-grained Atomic Evaluation of Factual Precision in Long Form Text Generation[C]//EMNLP. 2023.",
        "[11] CEMRI M C, PAN M Z, YANG S, et al. Why Do Multi-Agent LLM Systems Fail?[J]. arXiv:2503.13657, 2025.",
        "[12] LEE D, TIWARI M. Prompt Infection: LLM-to-LLM Prompt Injection within Multi-Agent Systems[J]. arXiv:2410.07283, 2024.",
        "[13] WANG S, ZHANG G, YU M, et al. G-Safeguard: A Topology-Guided Security Lens and Treatment on LLM-based Multi-agent Systems[C]//ACL. 2025.",
        "[14] 国务院. 新一代人工智能发展规划: 国发〔2017〕35号[Z]. 2017.",
        "[15] 国家互联网信息办公室, 国家发展和改革委员会, 教育部, 等. 生成式人工智能服务管理暂行办法[Z]. 2023.",
        "[16] 中华人民共和国网络安全法[Z]. 2016.",
        "[17] 中华人民共和国数据安全法[Z]. 2021.",
        "[18] 中华人民共和国个人信息保护法[Z]. 2021.",
        "[19] 国务院. 网络数据安全管理条例: 国务院令第790号[Z]. 2024.",
    ]
    for reference in references:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.hanging_indent = Pt(24)
        set_run_font(p.add_run(reference), size=10.5)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(DOCX_PATH)


def main() -> None:
    """生成图示和最终 DOCX。"""
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    configure_matplotlib()
    build_report(figure_ipjg(), figure_iptcm())
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
