#!/usr/bin/env python3
"""Build the implementation principle DOCX from markdown."""

from __future__ import annotations

import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".docx_deps"))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SRC = ROOT / "docs" / "implementation_principle_2_2_2_3.md"
OUT = ROOT / "docs" / "implementation_principle_2_2_to_2_5_4.docx"


def set_run_font(run, size: int = 12, code_style: bool = False) -> None:
    run.font.name = "Consolas" if code_style else "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str, code_style: bool = False, formula_style: bool = False) -> None:
    paragraph = doc.add_paragraph(text)
    if formula_style:
        paragraph.alignment = 1
        paragraph.paragraph_format.line_spacing = 1.2
    elif code_style:
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.1
    else:
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size=11 if formula_style else (10 if code_style else 12), code_style=code_style)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.size = Pt(size)


def main() -> None:
    doc = Document()
    configure_styles(doc)
    base = SRC.parent
    in_code = False
    in_formula = False
    code_lines: list[str] = []
    formula_lines: list[str] = []

    for raw_line in SRC.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.strip() == "$$":
            if in_formula:
                add_paragraph(doc, "\n".join(formula_lines), formula_style=True)
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
            continue
        if in_formula:
            formula_lines.append(line)
            continue
        if line.startswith("```"):
            if in_code:
                add_paragraph(doc, "\n".join(code_lines), code_style=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            caption = line[2 : line.index("](")]
            rel_path = line[line.index("](") + 2 : -1]
            image_path = base / rel_path
            paragraph = doc.add_paragraph()
            paragraph.alignment = 1
            paragraph.add_run().add_picture(str(image_path), width=Cm(15.5))
            caption_paragraph = doc.add_paragraph(caption)
            caption_paragraph.alignment = 1
            for run in caption_paragraph.runs:
                set_run_font(run, size=10)
        else:
            add_paragraph(doc, line)

    if code_lines:
        add_paragraph(doc, "\n".join(code_lines), code_style=True)
    if formula_lines:
        add_paragraph(doc, "\n".join(formula_lines), formula_style=True)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
