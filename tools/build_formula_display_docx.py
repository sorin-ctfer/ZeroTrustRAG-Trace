#!/usr/bin/env python3
"""Build a DOCX where LaTeX-style formulas are rendered as images."""

from __future__ import annotations

import re
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / ".docx_deps"))

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


SRC = ROOT / "docs" / "implementation_principle_2_2_2_3.md"
OUT = ROOT / "docs" / "implementation_principle_2_2_to_2_5_4_formula.docx"
FORMULA_DIR = ROOT / "docs" / "implementation_assets" / "formula_images"


def font_path() -> str:
    candidates = [
        "/mnt/c/Windows/Fonts/msyh.ttc",
        "/mnt/c/Windows/Fonts/simhei.ttf",
        "/mnt/c/Windows/Fonts/consola.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return path
    raise FileNotFoundError("No usable font found for formula rendering.")


FONT_PATH = font_path()


def latex_to_display_text(formula: str) -> str:
    """Convert common LaTeX commands to readable Unicode text for robust DOCX display."""
    text = formula.strip()
    replacements = {
        r"\mathcal{R}": "𝓡",
        r"\mathcal{D}": "𝓓",
        r"\mathcal{C}": "𝓒",
        r"\mathcal{M}": "𝓜",
        r"\mathcal{V}": "𝓥",
        r"\mathcal{I}": "𝓘",
        r"\mathcal{E}": "𝓔",
        r"\mathcal{G}": "𝓖",
        r"\mathbf{1}": "I",
        r"\cdot": "·",
        r"\times": "×",
        r"\wedge": "∧",
        r"\vee": "∨",
        r"\implies": "⇒",
        r"\leftarrow": "←",
        r"\rightarrow": "→",
        r"\xrightarrow": "→",
        r"\infty": "∞",
        r"\varepsilon": "ε",
        r"\theta": "θ",
        r"\lambda": "λ",
        r"\alpha": "α",
        r"\beta": "β",
        r"\gamma": "γ",
        r"\sum": "Σ",
        r"\prod": "Π",
        r"\cup": "∪",
        r"\cap": "∩",
        r"\emptyset": "∅",
        r"\forall": "∀",
        r"\exists": "∃",
        r"\neq": "≠",
        r"\ge": "≥",
        r"\le": "≤",
        r"\to": "→",
        r"\arg": "arg",
        r"\max": "max",
        r"\min": "min",
        r"\text": "",
        r"\begin{aligned}": "",
        r"\end{aligned}": "",
        r"\begin{cases}": "",
        r"\end{cases}": "",
        r"\begin{bmatrix}": "[",
        r"\end{bmatrix}": "]",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    text = re.sub(r"\\tag\{([^}]+)\}", r"    (\1)", text)
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", text)
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    text = text.replace("{", "").replace("}", "")
    text = text.replace("&=", "=").replace("&", "")
    text = text.replace("\\\\", "\n")
    text = re.sub(r"\s+\n", "\n", text)
    text = re.sub(r"\n\s+", "\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def render_formula(formula: str, index: int) -> Path:
    display = latex_to_display_text(formula)
    lines = display.splitlines() or [display]
    font = ImageFont.truetype(FONT_PATH, 30)
    small_font = ImageFont.truetype(FONT_PATH, 26)
    padding_x = 42
    padding_y = 26
    line_gap = 14
    dummy = Image.new("RGB", (10, 10), "white")
    draw = ImageDraw.Draw(dummy)
    widths = []
    heights = []
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    width = max(max(widths) + padding_x * 2, 900)
    height = sum(heights) + line_gap * (len(lines) - 1) + padding_y * 2
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    y = padding_y
    for line, line_height, line_width in zip(lines, heights, widths):
        x = (width - line_width) / 2
        draw.text((x, y), line, fill="black", font=font if len(line) < 120 else small_font)
        y += line_height + line_gap
    out = FORMULA_DIR / f"formula_{index:03d}.png"
    image.save(out)
    return out


def set_run_font(run, size: int = 12, code_style: bool = False) -> None:
    run.font.name = "Consolas" if code_style else "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def add_paragraph(doc: Document, text: str, code_style: bool = False) -> None:
    paragraph = doc.add_paragraph(text)
    if code_style:
        paragraph.paragraph_format.left_indent = Pt(18)
        paragraph.paragraph_format.line_spacing = 1.05
    else:
        paragraph.paragraph_format.first_line_indent = Pt(24)
        paragraph.paragraph_format.line_spacing = 1.5
    for run in paragraph.runs:
        set_run_font(run, size=10 if code_style else 12, code_style=code_style)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(12)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 13), ("Heading 4", 12)]:
        styles[style_name].font.name = "Times New Roman"
        styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[style_name].font.size = Pt(size)


def add_image(doc: Document, image_path: Path, caption: str | None = None, width_cm: float = 15.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = 1
    paragraph.add_run().add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        caption_paragraph = doc.add_paragraph(caption)
        caption_paragraph.alignment = 1
        for run in caption_paragraph.runs:
            set_run_font(run, size=10)


def is_formula_like_code(lines: list[str]) -> bool:
    text = "\n".join(lines).strip()
    if not text or len(lines) > 8:
        return False
    if text.startswith(("算法：", "输入：", "输出：", "步骤：", "Report = {", "Evidence = {")):
        return False
    math_tokens = ["=", "≥", "≤", "⇒", "∈", "∪", "∩", "θ", "λ", "α", "β", "γ", "×", "∞", "TopK", "sim("]
    return any(token in text for token in math_tokens)


def main() -> None:
    FORMULA_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    base = SRC.parent
    in_code = False
    in_formula = False
    code_lines: list[str] = []
    formula_lines: list[str] = []
    formula_index = 1

    for raw_line in SRC.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.strip() == "$$":
            if in_formula:
                image_path = render_formula("\n".join(formula_lines), formula_index)
                add_image(doc, image_path, width_cm=14.5)
                formula_index += 1
                formula_lines = []
                in_formula = False
            else:
                in_formula = True
            continue
        if in_formula:
            formula_lines.append(line)
            continue

        if line.startswith("```"):
            if in_code:
                if is_formula_like_code(code_lines):
                    image_path = render_formula("\n".join(code_lines), formula_index)
                    add_image(doc, image_path, width_cm=14.5)
                    formula_index += 1
                else:
                    add_paragraph(doc, "\n".join(code_lines), code_style=True)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if not line.strip():
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("![") and "](" in line and line.endswith(")"):
            caption = line[2 : line.index("](")]
            rel_path = line[line.index("](") + 2 : -1]
            add_image(doc, base / rel_path, caption=caption)
        elif line.startswith("---"):
            doc.add_paragraph("")
        else:
            add_paragraph(doc, line)

    if code_lines:
        if is_formula_like_code(code_lines):
            image_path = render_formula("\n".join(code_lines), formula_index)
            add_image(doc, image_path, width_cm=14.5)
        else:
            add_paragraph(doc, "\n".join(code_lines), code_style=True)
    if formula_lines:
        image_path = render_formula("\n".join(formula_lines), formula_index)
        add_image(doc, image_path, width_cm=14.5)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
